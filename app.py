import streamlit as st
import os
import requests
import re
import json
import time
import datetime
import base64
import hashlib
from fpdf import FPDF
from openai import OpenAI
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns
from io import BytesIO
from collections import Counter
import PyPDF2  # Libreria necessaria per leggere i PDF caricati
from PIL import Image

# ======================================================================================================================
# 0. GESTIONE MEMORIA DI STATO E PREVENZIONE AUTO-RESET
# ======================================================================================================================
# Questo blocco garantisce che l'applicazione mantenga i dati in memoria durante le elaborazioni lunghe
# e i cambi di tab. I dati verranno azzerati SOLO tramite l'esplicito pulsante di RESET.
if "memoria_blindata" not in st.session_state:
    st.session_state["memoria_blindata"] = True
    st.session_state["indice_raw"] = ""
    st.session_state["lista_capitoli"] = []
    st.session_state["conoscenza_extra"] = ""
    st.session_state["immagini_capitoli"] = {}

# Compatibilità con sessioni aperte prima dell'introduzione delle immagini.
if "immagini_capitoli" not in st.session_state:
    st.session_state["immagini_capitoli"] = {}

# ======================================================================================================================
# FUNZIONI DI SUPPORTO PER ANALISI DOCUMENTI (NUOVO MODULO)
# ======================================================================================================================
def estrai_testo_da_files(caricati):
    testo_totale = ""
    for file in caricati:
        try:
            if file.name.lower().endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    testo_totale += (page.extract_text() or "") + "\n"
            elif file.name.lower().endswith('.docx'):
                doc = Document(file)
                for para in doc.paragraphs:
                    testo_totale += para.text + "\n"
        except Exception as e:
            st.error(f"Errore nella lettura di {file.name}: {e}")
    return testo_totale

# ======================================================================================================================
# 1. ARCHITETTURA DI SISTEMA E SICUREZZA API
# ======================================================================================================================
# Nome Applicazione: AI di Antonino: Ebook Mondiale Creator PRO
# Developer: Antonino & Gemini Collaboration
# Core Update: Integrazione Neuromarketing (Triune Brain Methodology) con Motore Decisionale Dinamico.

# --- AGGIORNAMENTO SICUREZZA API ---
try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception as e:
    st.error("ERRORE CRITICO: Chiave API OpenAI non trovata nei Secrets di Streamlit. Assicurati di aver creato il file secrets.toml o configurato i Secrets online.")

st.set_page_config(
    page_title="AI di Antonino: Ebook Mondiale Creator PRO",
    layout="wide",
    initial_sidebar_state="expanded",
    page_icon="✒️"
)

# ======================================================================================================================
# 2. DIZIONARIO MULTILINGUA INTEGRALE (9 LINGUE GLOBALI - ESPANSO)
# ======================================================================================================================
TRADUZIONI = {
    "Italiano": {
        "side_tit": "⚙️ Configurazione Editor",
        "lbl_tit": "Titolo del Libro", "lbl_auth": "Nome Autore", "lbl_lang": "Lingua", 
        "lbl_gen": "Genere Letterario", "lbl_style": "Tipologia Scrittura", "lbl_plot": "Trama o Argomento",
        "lbl_narrative": "Stile di Racconto", "lbl_goal": "Obiettivo del Libro", "lbl_pov": "Punto di Vista (Pronome)",
        "btn_res": "🔄 RESET PROGETTO", "tabs": ["📊 1. Indice", "✍️ 2. Scrittura & Quiz", "📖 3. Anteprima", "📑 4. Esporta"],
        "btn_idx": "🚀 Genera Indice Professionale", "btn_sync": "✅ Salva e Sincronizza Capitoli",
        "lbl_sec": "Seleziona sezione:", "btn_write": "✨ SCRIVI CONTENUTO (Dettagliato)",
        "btn_quiz": "🧠 AGGIUNGI QUIZ AL LIBRO", "btn_edit": "🚀 RIELABORA CON IA",
        "msg_run": "Il neuro-linguista sta analizzando gerarchia, stile e target emotivo...", "preface": "Prefazione", "ack": "Ringraziamenti",
        "preview_tit": "📖 Vista Lettura Professionale", "btn_word": "📥 Scarica Word (.docx)", "btn_pdf": "📥 Scarica PDF (.pdf)",
        "msg_err_idx": "Genera l'indice nella Tab 1 prima di procedere.", "msg_success_sync": "Capitoli sincronizzati!",
        "label_editor": "Editor di Testo Professionale", "welcome": "👋 Benvenuto nell'Ebook Creator di Antonino.",
        "guide": "Usa la sidebar a sinistra per impostare i parametri del tuo libro."
    },
    "English": {
        "side_tit": "⚙️ Editor Setup", "lbl_tit": "Book Title", "lbl_auth": "Author Name", "lbl_lang": "Language", 
        "lbl_gen": "Genre", "lbl_style": "Writing Style", "lbl_plot": "Plot", "lbl_narrative": "Narrative Style", "lbl_goal": "Book Goal", "lbl_pov": "Point of View (Pronoun)",
        "btn_res": "🔄 RESET PROJECT", "tabs": ["📊 1. Index", "✍️ 2. Write & Quiz", "📖 3. Preview", "📑 4. Export"],
        "btn_idx": "🚀 Generate Index", "btn_sync": "✅ Sync Chapters", "lbl_sec": "Select section:",
        "btn_write": "✨ WRITE CONTENT", "btn_quiz": "🧠 ADD QUIZ", "btn_edit": "🚀 REWRITE",
        "msg_run": "Native expert analyzing hierarchy, style and goal...", "preface": "Preface", "ack": "Acknowledgements",
        "preview_tit": "📖 Reading View", "btn_word": "📥 Word", "btn_pdf": "📥 PDF",
        "msg_err_idx": "Generate index first.", "msg_success_sync": "Synced!",
        "label_editor": "Editor", "welcome": "👋 Welcome.", "guide": "Use sidebar."
    },
    "Español": {
        "side_tit": "⚙️ Configuración del Editor", "lbl_tit": "Título del Libro", "lbl_auth": "Nombre del Autor", "lbl_lang": "Idioma", 
        "lbl_gen": "Género Literario", "lbl_style": "Estilo de Escritura", "lbl_plot": "Trama o Argumento", "lbl_narrative": "Estilo Narrativo", "lbl_goal": "Objetivo del Libro", "lbl_pov": "Punto de Vista (Pronombre)",
        "btn_res": "🔄 RESETEAR PROYECTO", "tabs": ["📊 1. Índice", "✍️ 2. Escritura y Quiz", "📖 3. Vista Previa", "📑 4. Exportar"],
        "btn_idx": "🚀 Generar Índice Profesional", "btn_sync": "✅ Guardar y Sincronizar", "lbl_sec": "Seleccionar sección:",
        "btn_write": "✨ ESCRIBIR CONTENIDO", "btn_quiz": "🧠 AÑADIR QUIZ", "btn_edit": "🚀 REESCRIBIR",
        "msg_run": "Analizando jerarquía y estilo...", "preface": "Prefacio", "ack": "Agradecimientos",
        "preview_tit": "📖 Vista de Lectura", "btn_word": "📥 Descargar Word", "btn_pdf": "📥 Descargar PDF",
        "msg_err_idx": "Genera el índice primero.", "msg_success_sync": "¡Sincronizado!", "label_editor": "Editor Profesional", "welcome": "👋 Bienvenido.", "guide": "Usa la barra lateral."
    },
    "Français": {
        "side_tit": "⚙️ Configuration de l'Éditeur", "lbl_tit": "Titre du Livre", "lbl_auth": "Nom de l'Auteur", "lbl_lang": "Langue", 
        "lbl_gen": "Genre Littéraire", "lbl_style": "Style d'Écriture", "lbl_plot": "Intrigue ou Sujet", "lbl_narrative": "Style Narratif", "lbl_goal": "Objectif du Livre", "lbl_pov": "Point de Vue (Pronom)",
        "btn_res": "🔄 RÉINITIALISER", "tabs": ["📊 1. Index", "✍️ 2. Écriture & Quiz", "📖 3. Aperçu", "📑 4. Exporter"],
        "btn_idx": "🚀 Générer l'Index", "btn_sync": "✅ Synchroniser", "lbl_sec": "Sélectionner la section:",
        "btn_write": "✨ ÉCRIRE LE CONTENU", "btn_quiz": "🧠 AJOUTER UN QUIZ", "btn_edit": "🚀 RÉÉCRIRE",
        "msg_run": "Analyse de la hiérarchie et du style...", "preface": "Préface", "ack": "Remerciements",
        "preview_tit": "📖 Aperçu de Lecture", "btn_word": "📥 Télécharger Word", "btn_pdf": "📥 Télécharger PDF",
        "msg_err_idx": "Générez l'index d'abord.", "msg_success_sync": "Synchronisé!", "label_editor": "Éditeur Professionnel", "welcome": "👋 Bienvenue.", "guide": "Utilisez la barre latérale."
    },
    "Deutsch": {
        "side_tit": "⚙️ Editor-Setup", "lbl_tit": "Buchtitel", "lbl_auth": "Autorenname", "lbl_lang": "Sprache", 
        "lbl_gen": "Genre", "lbl_style": "Schreibstil", "lbl_plot": "Handlung", "lbl_narrative": "Erzählstil", "lbl_goal": "Buchziel", "lbl_pov": "Erzählperspektive (Pronomen)",
        "btn_res": "🔄 PROJEKT ZURÜCKSETZEN", "tabs": ["📊 1. Index", "✍️ 2. Schreiben & Quiz", "📖 3. Vorschau", "📑 4. Exportieren"],
        "btn_idx": "🚀 Index Generieren", "btn_sync": "✅ Synchronisieren", "lbl_sec": "Abschnitt wählen:",
        "btn_write": "✨ INHALT SCHREIBEN", "btn_quiz": "🧠 QUIZ HINZUFÜGEN", "btn_edit": "🚀 UMSCHREIBEN",
        "msg_run": "Analysiere Hierarchie und Stil...", "preface": "Vorwort", "ack": "Danksagungen",
        "preview_tit": "📖 Leseansicht", "btn_word": "📥 Word Herunterladen", "btn_pdf": "📥 PDF Herunterladen",
        "msg_err_idx": "Generiere zuerst den Index.", "msg_success_sync": "Synchronisiert!", "label_editor": "Professioneller Editor", "welcome": "👋 Willkommen.", "guide": "Nutze die Seitenleiste."
    },
    "Română": {
        "side_tit": "⚙️ Configurare Editor", "lbl_tit": "Titlul Cărții", "lbl_auth": "Nume Autor", "lbl_lang": "Limbă", 
        "lbl_gen": "Gen Literar", "lbl_style": "Stil de Scriere", "lbl_plot": "Subiect", "lbl_narrative": "Stil Narativ", "lbl_goal": "Obiectivul Cărții", "lbl_pov": "Punct de Vedere (Pronume)",
        "btn_res": "🔄 RESETARE PROIECT", "tabs": ["📊 1. Cuprins", "✍️ 2. Scriere & Quiz", "📖 3. Previzualizare", "📑 4. Export"],
        "btn_idx": "🚀 Generare Cuprins", "btn_sync": "✅ Sincronizare", "lbl_sec": "Selectează secțiunea:",
        "btn_write": "✨ SCRIE CONȚINUT", "btn_quiz": "🧠 ADAUGĂ QUIZ", "btn_edit": "🚀 RESCRIE",
        "msg_run": "Se analizează ierarhia și stilul...", "preface": "Prefață", "ack": "Mulțumiri",
        "preview_tit": "📖 Mod Citire", "btn_word": "📥 Descarcă Word", "btn_pdf": "📥 Descarcă PDF",
        "msg_err_idx": "Generează cuprinsul mai întâi.", "msg_success_sync": "Sincronizat!", "label_editor": "Editor Profesional", "welcome": "👋 Bun venit.", "guide": "Folosește bara lateral."
    },
    "Русский": {
        "side_tit": "⚙️ Настройки Редактора", "lbl_tit": "Название Книги", "lbl_auth": "Имя Автора", "lbl_lang": "Язык", 
        "lbl_gen": "Жанр", "lbl_style": "Стиль Написания", "lbl_plot": "Сюжет", "lbl_narrative": "Стиль Повествования", "lbl_goal": "Цель Книги", "lbl_pov": "Точка зрения (Местоимение)",
        "btn_res": "🔄 СБРОСИТЬ ПРОЕКТ", "tabs": ["📊 1. Оглавление", "✍️ 2. Текст и Тест", "📖 3. Просмотр", "📑 4. Export"],
        "btn_idx": "🚀 Создать Оглавление", "btn_sync": "✅ Синхронизировать", "lbl_sec": "Выберите раздел:",
        "btn_write": "✨ НАПИСАТЬ ТЕКСТ", "btn_quiz": "🧠 ДОБАВИТЬ ТЕСТ", "btn_edit": "🚀 ПЕРЕПИСАТЬ",
        "msg_run": "Анализ иерархии и стиля...", "preface": "Предисловие", "ack": "Благодарности",
        "preview_tit": "📖 Режим Чтения", "btn_word": "📥 Скачать Word", "btn_pdf": "📥 Скачать PDF",
        "msg_err_idx": "Сначала создайте оглавление.", "msg_success_sync": "Синхронизировано!", "label_editor": "Профессиональный Редактор", "welcome": "👋 Добро пожаловать.", "guide": "Используйте боковую панель."
    },
    "العربية": {
        "side_tit": "⚙️ إعدادات المحرر", "lbl_tit": "عنوان الكتاب", "lbl_auth": "اسم المؤلف", "lbl_lang": "اللغة", 
        "lbl_gen": "النوع الأدبي", "lbl_style": "أسلوب الكتابة", "lbl_plot": "الحبكة أو الموضوع", "lbl_narrative": "الأسلوب السردي", "lbl_goal": "هدف الكتاب", "lbl_pov": "وجهة النظر (الضمير)",
        "btn_res": "🔄 إعادة ضبط المشروع", "tabs": ["📊 1. الفهرس", "✍️ 2. الكتابة والاختبار", "📖 3. معاينة", "📑 4. تصدير"],
        "btn_idx": "🚀 إنشاء فهرس احترافي", "btn_sync": "✅ حفظ ومزامنة الفصول", "lbl_sec": "اختر القسم:",
        "btn_write": "✨ كتابة المحتوى", "btn_quiz": "🧠 إضافة اختبار", "btn_edit": "🚀 إعادة صياغة",
        "msg_run": "جاري تحليل التسلسل الهرمي والأسلوب...", "preface": "مقدمة", "ack": "شكر وتقدير",
        "preview_tit": "📖 عرض القراءة الاحترافي", "btn_word": "📥 تحميل Word", "btn_pdf": "📥 تحميل PDF",
        "msg_err_idx": "قم بإنشاء الفهرس أولاً.", "msg_success_sync": "تمت المزامنة!", "label_editor": "محرر نصوص احترافي", "welcome": "👋 مرحباً بك.", "guide": "استخدم الشريط الجانبي."
    },
    "中文": {
        "side_tit": "⚙️ 编辑器设置", "lbl_tit": "书名", "lbl_auth": "作者姓名", "lbl_lang": "语言", 
        "lbl_gen": "文学体裁", "lbl_style": "写作类型", "lbl_plot": "情节或主题", "lbl_narrative": "叙事风格", "lbl_goal": "书籍目标", "lbl_pov": "叙事视角 (代词)",
        "btn_res": "🔄 重置项目", "tabs": ["📊 1. 目录", "✍️ 2. 写作与测试", "📖 3. 预览", "📑 4. 导出"],
        "btn_idx": "🚀 生成专业目录", "btn_sync": "✅ 保存并同步章节", "lbl_sec": "选择章节:",
        "btn_write": "✨ 编写内容", "btn_quiz": "🧠 添加测试", "btn_edit": "🚀 用AI重写",
        "msg_run": "正在分析层级、风格和情感目标...", "preface": "前言", "ack": "致谢",
        "preview_tit": "📖 专业阅读视图", "btn_word": "📥 下载 Word", "btn_pdf": "📥 下载 PDF",
        "msg_err_idx": "请先生成目录。", "msg_success_sync": "已同步！", "label_editor": "专业文本编辑器", "welcome": "👋 欢迎。", "guide": "请使用左侧边栏设置书籍参数。"
    }
}

# ======================================================================================================================
# 3. BLOCCO CSS: SIDEBAR SCURA E PULSANTI SCURI (FORZATURA !IMPORTANT)
# ======================================================================================================================
st.markdown("""
<style>
#MainMenu, footer, header, [data-testid="stHeader"] {visibility: hidden;}
[data-testid="collapsedControl"] { display: none !important; }

section[data-testid="stSidebar"] { 
    min-width: 420px !important; max-width: 420px !important; 
    width: 420px !important; display: block !important; visibility: visible !important;
    transform: none !important; background-color: #1e1e1e !important; border-right: 1px solid #333;
}
/* Sidebar fissa: impedisce il collasso accidentale della barra laterale. */
section[data-testid="stSidebar"][aria-expanded="false"],
[data-testid="stSidebar"][aria-expanded="false"] {
    min-width: 420px !important; max-width: 420px !important; width: 420px !important;
    display: block !important; visibility: visible !important; transform: none !important;
}
section[data-testid="stSidebar"] > div:first-child {
    width: 420px !important; min-width: 420px !important; display: block !important;
}
section[data-testid="stSidebar"] .stMarkdown, section[data-testid="stSidebar"] label, 
section[data-testid="stSidebar"] p, section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
}
.stButton>button {
    width: 100% !important; border-radius: 10px !important; height: 4.2em !important; 
    font-weight: bold !important; background-color: #1e1e1e !important; color: #ffffff !important;
    font-size: 18px !important; border: 2px solid #333 !important; 
    box-shadow: 0px 4px 12px rgba(0, 0, 0, 0.4) !important; transition: all 0.3s ease !important;
}
.stButton>button:hover { 
    background-color: #333333 !important; border-color: #007BFF !important; 
    color: #007BFF !important; transform: translateY(-2px) !important;
}
.preview-box {
    background-color: #ffffff !important; padding: 80px; border: 1px solid #ccc; 
    border-radius: 4px; height: 900px; overflow-y: scroll;
    font-family: 'Times New Roman', serif; line-height: 2.0; 
    color: #111 !important; box-shadow: 0px 25px 60px rgba(0,0,0,0.2); margin: 0 auto;
}
.custom-title {
    font-size: 38px; font-weight: 900; color: #111; text-align: center;
    padding: 30px; background-color: #ffffff; border-radius: 12px;
    margin-bottom: 30px; border-bottom: 6px solid #1e1e1e;
}
div[data-baseweb="select"] > div { background-color: #2b2b2b !important; color: white !important; }
</style>
""", unsafe_allow_html=True)

# ======================================================================================================================
# 4. GESTIONE EXPORT PDF (CHIRURGIA: FIX TITOLI LUNGHI E MARGINI)
# ======================================================================================================================
class EbookPDF(FPDF):
    def __init__(self, titolo, autore):
        super().__init__()
        self.titolo = self._clean(titolo)
        self.autore = self._clean(autore)
        
        # --- FIX MARGINI: Imposta margini espliciti e interruzione pagina automatica ---
        # Imposta margine sinistro, superiore e destro a 15 mm
        self.set_margins(15, 15, 15)
        # Forza il salto pagina automatico quando si arriva a 15 mm dal fondo
        self.set_auto_page_break(auto=True, margin=15)
        
    def _clean(self, txt):
        """Sanitizzazione forzata per FPDF latin-1. Evita crash da smart quotes e unicode."""
        if not txt: return ""
        replacements = {'“': '"', '”': '"', '‘': "'", '’': "'", '—': '-', '–': '-', '…': '...'}
        for k, v in replacements.items(): 
            txt = txt.replace(k, v)
        return txt.encode('latin-1', 'replace').decode('latin-1')

    def header(self):
        if self.page_no() > 1:
            self.set_font('Arial', 'I', 9); self.set_text_color(150)
            self.cell(0, 10, f"{self.titolo} - {self.autore}", 0, 0, 'R'); self.ln(15)
            
    def footer(self):
        if self.page_no() <= 1:
            return
        self.set_y(-20); self.set_font('Arial', 'I', 9)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
        
    def cover_page(self):
        self.add_page(); self.set_font('Arial', 'B', 32); self.ln(100)
        self.multi_cell(0, 15, self.titolo.upper(), 0, 'C'); self.ln(20)
        self.set_font('Arial', 'I', 20); self.cell(0, 10, f"di {self.autore}", 0, 1, 'C')
        
    def add_content(self, title, content, image_bytes=None, image_caption=None):
        self.add_page(); self.ln(15); self.set_font('Arial', 'B', 22)
        # FIX: Sostituito cell() con multi_cell() per il titolo, per mandare a capo i titoli lunghi!
        self.multi_cell(0, 15, self._clean(title).upper(), 0, 'L'); self.ln(10); self.set_font('Arial', '', 12)
        # multi_cell con w=0 ora calcola la larghezza rispettando il margine destro (15mm)
        if image_bytes:
            image_path = os.path.join(st.session_state.get("tmp_dir", os.path.abspath("tmp")), "ebook_creator_image.png")
            os.makedirs(os.path.dirname(image_path), exist_ok=True)
            with open(image_path, "wb") as f:
                f.write(image_bytes)
            try:
                # Immagine compatta: lascia spazio alla prosa e mantiene proporzioni corrette.
                self.image(image_path, x=48, w=115)
                self.ln(4)
                if image_caption:
                    self.set_font('Arial', 'I', 9)
                    self.multi_cell(0, 6, self._clean(image_caption), 0, 'C')
                    self.ln(5)
            finally:
                try: os.remove(image_path)
                except OSError: pass
        self.set_font('Arial', '', 12)
        self.multi_cell(0, 10, self._clean(content))

# ======================================================================================================================
# 5. CORE LOGIC GPT-4o & ANALISI QUALITÀ (POTENZIATA) E DECISIONE NEURALE
# ======================================================================================================================
def chiedi_gpt(prompt, system_prompt):
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini", 
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
            temperature=0.75 
        )
        testo = response.choices[0].message.content.strip()
        prefissi = ["ecco", "certamente", "sicuramente", "ok", "here is", "sure"]
        righe = [l for l in testo.split("\n") if not any(l.lower().startswith(p) for p in prefissi)]
        return "\n".join(righe).strip()
    except Exception as e: return f"ERRORE: {str(e)}"

def verifica_e_correggi_fatti_online(testo, sezione, lingua):
    """Verifica soltanto i fatti aggiornabili che meritano una ricerca online."""
    try:
        risposta = client.responses.create(
            model="gpt-5-mini",
            tools=[{"type": "web_search_preview"}],
            input=(
                f"Verifica il testo seguente in lingua {lingua} relativo alla sezione '{sezione}'. "
                "Cerca online fonti autorevoli e aggiornate per ogni fatto verificabile, soprattutto "
                "leggi, normative, prezzi, licenze, specifiche, date, software e dati numerici. "
                "Correggi soltanto le affermazioni non aggiornate o non supportate; non inventare dati. "
                "Mantieni struttura e stile, ma NON inserire nel testo URL, link Markdown, citazioni, "
                "note bibliografiche, nomi di fonti o una sezione 'Fonti verificate'. Le fonti servono "
                "esclusivamente per il controllo interno e non devono comparire nell'opera destinata al lettore. "
                "Distingui i fatti verificati dagli esempi ipotetici senza apporre etichette tecniche o note di fonte. "
                "Restituisci solo il testo editoriale revisionato e pulito.\n\n"
                f"TESTO:\n{testo}"
            )
        )
        return pulisci_testo_editoriale(getattr(risposta, "output_text", None) or testo)
    except Exception as e:
        st.warning(f"Verifica online non disponibile: {e}")
        return pulisci_testo_editoriale(testo)


def richiede_verifica_fatti(testo, sezione=""):
    """Evita ricerche inutili per scene, esercizi e spiegazioni stabili."""
    campione = f"{sezione}\n{testo}".lower()
    indicatori = (
        "legge", "normativa", "regolamento", "decreto", "licenza", "prezzo", "tariffa",
        "syllabus", "soglia di superamento", "punteggio minimo", "durata dell'esame",
        "versione", "requisiti di sistema", "compatibilità", "aggiornamento software",
        "aggiornata al", "in vigore", "20/", "€", "$"
    )
    return any(indicatore in campione for indicatore in indicatori)


def audit_fatti_capitolo(capitolo, contenuti, lingua):
    """Esegue un solo controllo online sul capitolo completo, senza riscrivere le singole sezioni."""
    testo = "\n\n".join(f"SEZIONE: {nome}\n{contenuto}" for nome, contenuto in contenuti if contenuto.strip())
    if not testo or not richiede_verifica_fatti(testo, capitolo):
        return "Controllo fatti del capitolo non necessario: nessun dato variabile rilevato."
    try:
        risposta = client.responses.create(
            model="gpt-5-mini",
            tools=[{"type": "web_search_preview"}],
            input=(
                f"Controlla i soli fatti aggiornabili nel capitolo '{capitolo}', in lingua {lingua}. "
                "Verifica esclusivamente regole, norme, date, soglie, prezzi, versioni software, licenze e specifiche. "
                "Non riscrivere il capitolo e non citare fonti o URL. Restituisci soltanto: "
                "ESITO: nessuna correzione necessaria oppure una lista di correzioni puntuali con sezione e formulazione da aggiornare.\n\n"
                f"CAPITOLO:\n{testo}"
            )
        )
        return pulisci_testo_editoriale(getattr(risposta, "output_text", "") or "Controllo non disponibile.")
    except Exception as e:
        return f"Controllo fatti del capitolo non disponibile: {e}"

def pulisci_testo_editoriale(testo):
    """Rimuove fonti tecniche dal testo destinato ad anteprima ed esportazione."""
    if not testo:
        return ""
    testo = str(testo)
    # Rimuove Markdown e segni di formattazione tecnica: l'editor impagina il testo in modo nativo.
    testo = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", testo)
    testo = testo.replace("**", "").replace("__", "")
    testo = re.sub(r"(?m)^\s*>\s?", "", testo)
    testo = re.sub(r"(?is)(?:^|\n)\s{0,3}(?:#+\s*)?(?:fonti verificate|fonti consultate|riferimenti bibliografici|sources|references)\s*:?.*$", "", testo)
    testo = re.sub(r"\[([^\]]+)\]\(https?://[^)]+\)", r"\1", testo)
    testo = re.sub(r"https?://[^\s)\]>]+", "", testo)
    testo = re.sub(r"\s*\([^\n()]{0,180}(?:\.com|\.org|\.gov|\.edu|doi\.org|utm_source|consultato il)[^\n()]*\)", "", testo, flags=re.I)
    testo = re.sub(r"(?im)^\s*\[?(?:informazione|fatto|esempio|fonte)[^\n]{0,120}(?:da verificare|verificato|ipotetico|di carattere generale)[^\n]*\]?\s*$", "", testo)
    testo = re.sub(r"(?m)^\s*[-_*]{3,}\s*$", "", testo)
    testo = re.sub(r"\n{3,}", "\n\n", testo)
    return testo.strip()

def genera_immagine_capitolo(sezione, titolo, genere, trama, contenuto, lingua):
    """GPT-4o-mini prepara il brief; GPT-Image-1 Mini genera il visual economico."""
    contesto_basso = f"{titolo} {trama} {sezione}".lower()
    if "fusion 360" in contesto_basso:
        vincoli_dominio = (
            "Per Fusion 360, se il testo descrive l'interfaccia, rappresenta barra degli strumenti in alto, "
            "browser a sinistra, area di lavoro 3D centrale, timeline in basso e pannello contestuale a destra. "
            "Se descrive una procedura CAD, mostra invece soltanto le fasi e gli oggetti realmente nominati."
        )
    elif any(x in contesto_basso for x in ("ricetta", "cucina", "ricettario")):
        vincoli_dominio = "Per ricette e cucina, mostra ingredienti, utensili e passaggi culinari realmente descritti, senza testo nell'immagine."
    elif any(x in contesto_basso for x in ("romanzo", "thriller", "fantasy", "rosa", "narrativo")):
        vincoli_dominio = "Per narrativa, mostra una scena coerente con luogo, personaggi, atmosfera e azione del brano, senza inserire eventi non presenti."
    elif any(x in contesto_basso for x in ("business", "marketing", "finanza", "economia")):
        vincoli_dominio = "Per business ed economia, mostra relazioni, flussi, strumenti o situazioni operative citate, senza numeri o dati inventati."
    else:
        vincoli_dominio = "Adatta la rappresentazione al dominio del sottocapitolo: mostra esclusivamente oggetti, persone, processi o relazioni realmente descritti nel testo."
    descrizione = chiedi_gpt(
        f"Analizza esclusivamente il sottocapitolo '{sezione}' del libro '{titolo}'. "
        f"Argomento generale: {trama}. Genere: {genere}. Lingua: {lingua}.\n\n"
        "Crea un brief visivo strutturato e concreto con queste voci: "
        "CONCETTO CENTRALE; ELEMENTI OBBLIGATORI (solo quelli realmente descritti); "
        "POSIZIONE E RELAZIONI SPAZIALI; AZIONE O PROCEDURA DA MOSTRARE; "
        "DETTAGLI TECNICI DA RENDERE VISIBILI; ELEMENTI DA ESCLUDERE. "
        "Ogni elemento dell'immagine deve corrispondere a un'informazione del testo. "
        "Non creare una schermata CAD generica e non inventare pannelli, icone o funzioni. "
        "Se il testo descrive un'interfaccia, rappresenta chiaramente le zone nominate "
        "(browser, area di modellazione, barra strumenti, pannello proprietà) nella posizione coerente. "
        "Se descrive una procedura, mostra le fasi in sequenza con forme e frecce non testuali. "
        "VIETATO inserire parole, titoli, paragrafi, numeri, etichette, didascalie, loghi "
        "o schermate con testo nell'immagine. Restituisci solo il brief.\n\n"
        f"{vincoli_dominio}\nContenuto già scritto: {contenuto[-2500:]}",
        "Sei un instructional designer tecnico: produci brief visivi accurati e verificabili."
    )
    try:
        risposta = client.images.generate(model="gpt-image-1-mini", prompt=f"Crea un'immagine didattica di alta qualità per il genere '{genere}' e il sottocapitolo '{sezione}'. Segui alla lettera questo brief visivo, senza aggiungere elementi non richiesti:\n{descrizione}\n\n{vincoli_dominio}\nLa scena deve avere corrispondenza uno-a-uno con il testo. Scegli composizione, livello di dettaglio e linguaggio visivo appropriati al dominio e al pubblico: diagramma o tavola tecnica per manuali, scena concreta per procedure, composizione narrativa per narrativa, visualizzazione concettuale per saggistica. Non creare immagini generiche o astratte e non inventare funzioni, dati, persone o oggetti. Nessun testo, lettera, numero, titolo, didascalia o logo nell'immagine. Mantieni sfondo bianco, tratto nero, scala di grigi e stile monocromatico pulito.", size="1024x1024", quality="medium")
        dato = risposta.data[0]
        raw = None
        if getattr(dato, "b64_json", None): raw = base64.b64decode(dato.b64_json)
        elif getattr(dato, "url", None): raw = requests.get(dato.url, timeout=60).content
        if raw:
            # Riduce risoluzione/peso e converte sempre in bianco e nero prima di salvare.
            img = Image.open(BytesIO(raw)).convert("L")
            img.thumbnail((600, 600), Image.Resampling.LANCZOS)
            out = BytesIO(); img.save(out, format="PNG", optimize=True)
            return out.getvalue(), descrizione
        raise ValueError("Risposta immagini priva di dati utilizzabili")
    except Exception as e:
        st.error(f"Errore nella generazione dell'immagine: {e}")
        return None, None

def normalizza_immagine_caricata(file_caricato):
    """Prepara un'immagine caricata dall'utente per anteprima, Word e PDF."""
    try:
        sorgente = Image.open(BytesIO(file_caricato.getvalue()))
        if sorgente.mode in ("RGBA", "LA"):
            sfondo = Image.new("RGB", sorgente.size, "white")
            sfondo.paste(sorgente, mask=sorgente.getchannel("A"))
            sorgente = sfondo
        else:
            sorgente = sorgente.convert("RGB")
        sorgente.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
        output = BytesIO()
        sorgente.save(output, format="PNG", optimize=True)
        return output.getvalue()
    except Exception as e:
        st.error(f"Il file caricato non è un'immagine valida: {e}")
        return None

def elimina_paragrafo_docx(paragrafo):
    elemento = paragrafo._element
    elemento.getparent().remove(elemento)
    paragrafo._p = paragrafo._element = None

def aggiungi_numeri_pagina_docx(documento):
    """Inserisce il campo numero pagina nel piè di pagina di ogni sezione Word."""
    for sezione in documento.sections:
        paragrafo = sezione.footer.paragraphs[0]
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        campo_inizio = OxmlElement('w:fldChar')
        campo_inizio.set(ns.qn('w:fldCharType'), 'begin')
        istruzione = OxmlElement('w:instrText')
        istruzione.text = 'PAGE'
        campo_fine = OxmlElement('w:fldChar')
        campo_fine.set(ns.qn('w:fldCharType'), 'end')
        run = paragrafo.add_run()
        run._r.append(campo_inizio)
        run._r.append(istruzione)
        run._r.append(campo_fine)

def formatta_manoscritto_kdp(file_docx):
    """Applica un formato Word pulito 6x9 per il manoscritto KDP caricato dall'utente."""
    documento = Document(BytesIO(file_docx.getvalue()))
    for nome_stile in ('Heading 1', 'Heading 2'):
        try:
            documento.styles[nome_stile]
        except KeyError:
            documento.styles.add_style(nome_stile, WD_STYLE_TYPE.PARAGRAPH)

    for sezione in documento.sections:
        sezione.page_width = Inches(6)
        sezione.page_height = Inches(9)
        sezione.top_margin = Inches(0.75)
        sezione.bottom_margin = Inches(0.75)
        sezione.left_margin = Inches(0.75)
        sezione.right_margin = Inches(0.75)

    for paragrafo in list(documento.paragraphs):
        testo = pulisci_testo_editoriale(paragrafo.text).strip()
        if not testo:
            elimina_paragrafo_docx(paragrafo)
            continue
        paragrafo.text = ' '.join(testo.split())
        if len(paragrafo.text) < 80 and re.search(r'(?i)\b(capitolo|chapter|parte|part)\b', paragrafo.text):
            paragrafo.style = 'Heading 1'
            paragrafo.paragraph_format.page_break_before = True
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragrafo.paragraph_format.space_before = Pt(0)
            paragrafo.paragraph_format.space_after = Pt(30)
        elif len(paragrafo.text) < 100 and re.match(r'^\d+(?:\.\d+)?\s+', paragrafo.text):
            paragrafo.style = 'Heading 2'
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragrafo.paragraph_format.first_line_indent = Inches(0)
            paragrafo.paragraph_format.space_before = Pt(18)
            paragrafo.paragraph_format.space_after = Pt(10)
        else:
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragrafo.paragraph_format.first_line_indent = Inches(0.25)
            paragrafo.paragraph_format.space_after = Pt(6)

    stile_normale = documento.styles['Normal']
    stile_normale.font.name = 'Georgia'
    stile_normale.font.size = Pt(11)
    aggiungi_numeri_pagina_docx(documento)
    output = BytesIO()
    documento.save(output)
    output.seek(0)
    return output

def estrai_anteprima_manoscritto(file_caricato):
    """Estrae una porzione di testo da DOCX o PDF per la generazione dei metadati."""
    dati = BytesIO(file_caricato.getvalue())
    if file_caricato.name.lower().endswith('.docx'):
        documento = Document(dati)
        return '\n'.join(p.text for p in documento.paragraphs[:100])
    lettore = PyPDF2.PdfReader(dati)
    return '\n'.join((pagina.extract_text() or '') for pagina in lettore.pages[:15])

def analizza_qualita_prosa(testo):
    """
    Motore Linter NLP Potenziato: analizza densità, lunghezza frasi e vocabolario.
    """
    if not testo or len(testo) < 50: 
        return "⚠️ Testo troppo breve per un'analisi sintattica significativa."
    
    risultati = ["📊 **REPORT LINTER AVANZATO E ANALISI SINTATTICA**\n"]
    
    # 1. Parsing base
    parole = re.findall(r'\b\w+\b', testo.lower())
    frasi = [f.strip() for f in re.split(r'[.!?]+', testo) if len(f.strip()) > 5]
    
    tot_parole = len(parole)
    tot_frasi = len(frasi) if len(frasi) > 0 else 1
    
    # 2. Diversità Lessicale (Ricchezza del vocabolario)
    vocabolo_unico = len(set(parole))
    indice_diversita = (vocabolo_unico / tot_parole) * 100 if tot_parole > 0 else 0
    if indice_diversita < 35:
        risultati.append(f"⚠️ **Vocabolario Ripetitivo**: Indice di diversità lessicale basso ({indice_diversita:.1f}%). Valuta di usare più sinonimi.")
    else:
        risultati.append(f"✅ **Ricchezza Lessicale**: Ottima diversità ({indice_diversita:.1f}%). Il testo risulta stimolante.")

    # 3. Lunghezza Media delle Frasi (Pacing e Affaticamento Neocorteccia)
    parole_per_frase = tot_parole / tot_frasi
    if parole_per_frase > 30:
        risultati.append(f"⚠️ **Sintassi Pesante**: Le frasi sono troppo lunghe (media {parole_per_frase:.1f} parole/frase). Rischio di affaticamento cognitivo: spezza i periodi.")
    elif parole_per_frase < 8:
        risultati.append(f"⚠️ **Ritmo Frammentato**: Frasi molto brevi (media {parole_per_frase:.1f} parole/frase). Il testo potrebbe risultare troppo robotico o telegrafico.")
    else:
        risultati.append(f"✅ **Ritmo e Leggibilità**: Lunghezza frasi perfettamente bilanciata (media {parole_per_frase:.1f} parole/frase).")

    # 4. Ripetizioni Ravvicinate Fastidiose (Finestra Mobile)
    ripetizioni = []
    for i in range(len(parole) - 15):
        target = parole[i]
        # Escludiamo congiunzioni e preposizioni comuni basandoci sulla lunghezza della parola
        if len(target) > 4 and target in parole[i+1 : i+15]: 
            ripetizioni.append(target)
            
    if ripetizioni:
        comuni = [p[0] for p in Counter(ripetizioni).most_common(5)]
        risultati.append(f"🔍 **Allerta Ripetizioni Ravvicinate**: Le seguenti parole si ripetono troppo vicine tra loro: *{', '.join(comuni)}*")
    else:
        risultati.append("✅ **Fluidità Testuale**: Nessuna ripetizione fastidiosa o eco ravvicinata rilevata.")

    return "\n\n".join(risultati)

def sync_capitoli():
    testo_indice = st.session_state.get("indice_raw", "")
    if not testo_indice: st.session_state['lista_capitoli'] = []; return
    lista = []
    
    # --- INIZIO RIGHE MODIFICATE PER SUPPORTARE TUTTE LE LINGUE NELL'ESTRAZIONE DELLA SIDEBAR ---
    # Inclusi tutti i termini stranieri di "Capitolo" o "Parte" per il Parsing
    regex = r'(?i)(Capitolo|Chapter|Kapitel|Capítulo|Chapitre|Capitolul|Глава|الفصل|Раздел|章节|Secţiune|Parte|Part|Partie|Teil|Partea|Часть|الجزء|部分|\d+\.)'
    # --- FINE RIGHE MODIFICATE ---
    
    for riga in testo_indice.split('\n'):
        if re.search(regex, riga.strip()): lista.append(riga.strip())
    st.session_state['lista_capitoli'] = lista


# ======================================================================================================================
# PROFILI EDITORIALI: REGOLE SPECIFICHE PER GENERE, TIPOLOGIA E STRUTTURA
# ======================================================================================================================
def profilo_tipologia_stesura(stile):
    """Restituisce istruzioni di stesura realmente diverse per ogni tipologia selezionabile."""
    profili = {
        "Standard": "Esponi con chiarezza e ordine. Alterna spiegazione, esempio e applicazione senza estremi retorici.",
        "Professionale Accademico": "Definisci termini, separa fatti, metodo, interpretazioni e limiti. Usa un registro preciso e prudente; non trasformare il testo in un elenco di istruzioni quando il contenuto richiede argomentazione.",
        "Persuasivo (Neuromarketing Applicato)": "Parti da un problema concreto, chiarisci valore e prove, affronta obiezioni e guida verso una scelta o un'azione. Non usare pressione, manipolazione o promesse garantite.",
        "Conversazionale ed Empatico": "Accompagna il lettore con un linguaggio umano e rispettoso. Anticipa dubbi reali, normalizza gli ostacoli e offri indicazioni applicabili senza toni paternalistici.",
        "Scientifico Divulgativo": "Rendi comprensibili concetti complessi attraverso definizioni semplici, meccanismi, esempi e limiti. Distingui sempre dati, ipotesi, analogie e aspetti da verificare.",
        "Storytelling Immersivo": "Costruisci scene, azioni, conseguenze e dettagli sensoriali coerenti. Ogni sezione deve far evolvere conflitto, personaggio, relazione o posta in gioco; non riassumere ciò che può essere mostrato.",
        "Giornalistico d'Inchiesta": "Mantieni una linea di verifica: fatti documentabili, fonti da controllare, contraddizioni, contesto e conseguenze. Non presentare ipotesi come prove e non inventare testimonianze.",
        "Socratico (Dialogico / Riflessivo)": "Organizza la sezione attorno a una domanda reale. Esplora presupposti, dubbi e obiezioni, quindi porta il lettore a una conclusione argomentata o a una riflessione verificabile.",
        "Epico ed Evocativo": "Usa immagini e ritmo evocativi senza perdere chiarezza. La trasformazione, le prove e il significato devono essere concreti e adeguati al genere, non formule decorative.",
        "Minimalista ed Essenziale": "Elimina tutto ciò che non serve. Usa frasi sobrie, titoli funzionali, esempi strettamente necessari e una sola idea centrale per blocco di testo."
    }
    return profili.get(stile, profili["Standard"])


def profilo_genere_stesura(genere):
    """Regole di forma e contenuto per tutti i generi offerti dall'interfaccia."""
    profili = {
        "Saggio Scientifico": "Sostieni una tesi con definizioni, metodo, evidenze, controargomentazioni, limiti e implicazioni. Non inventare dati o studi.",
        "Quiz Scientifico": "Alterna spiegazione essenziale, domande verificabili, soluzioni motivate e chiarimento degli errori più probabili.",
        "Manuale Tecnico": "Fornisci prerequisiti, strumenti, parametri, sequenze operative, controlli, errori e criteri di riuscita. Se software o norme possono cambiare, segnala cosa verificare.",
        "Religioso / Teologico": "Distingui testi, interpretazioni, tradizioni e opinioni. Mantieni rispetto, precisione storica e nessuna affermazione dogmatica non attribuita.",
        "Spirituale / Esoterico": "Usa un tono rispettoso e non prescrittivo. Presenta pratiche come esperienze personali o tradizionali, non come cure o certezze scientifiche.",
        "Meditazione / Mindfulness": "Offri pratiche graduali, istruzioni sicure, durata indicativa, osservazioni e alternative. Evita promesse terapeutiche o risultati garantiti.",
        "Business & Marketing": "Usa obiettivi, pubblico, casi, metriche, scelte operative e criteri di verifica. Se i dati non sono forniti, usa esempi dichiaratamente ipotetici.",
        "Economia e Finanza": "Separa educazione generale da consulenza personalizzata. Spiega rischio, limiti, dati e ipotesi; non dare raccomandazioni finanziarie individuali.",
        "Romanzo Rosa": "Sviluppa desiderio, relazione, vulnerabilità, ostacoli e scelta emotiva attraverso scene, dialoghi e trasformazione dei personaggi.",
        "Thriller / Noir": "Costruisci tensione con indizi, conseguenze, conflitti e rivelazioni coerenti. Ogni capitolo deve cambiare le informazioni disponibili o aumentare la posta in gioco.",
        "Fantasy": "Mantieni coerenti mondo, regole, conflitti e conseguenze. Mostra il worldbuilding dentro azioni e scene, senza blocchi enciclopedici.",
        "Fantascienza": "Rendi coerente la premessa speculativa e mostra come modifica società, tecnologia, personaggi e conflitto. Non sostituire la storia con spiegazioni astratte.",
        "Manuale Psicologico": "Spiega modelli e pratiche in modo accessibile, con limiti chiari. Non fare diagnosi, non promettere cura e invita a rivolgersi a professionisti quando necessario.",
        "Biografia": "Segui una cronologia significativa, usando fonti verificabili e distinguendo fatti, testimonianze e interpretazioni. Privilegia svolte e contesto rispetto a elenchi di date.",
        "Ricettario": "Ogni capitolo-ricetta deve contenere porzioni, tempi, ingredienti con dosi, procedimento numerato, segnali di riuscita, errore e correzione, variante e conservazione solo se verificata. Non duplicare la stessa ricetta in forma breve ed estesa.",
        "Test Prep (Preparazione Esami)": "Spiega soltanto le competenze pertinenti alla prova, poi fornisci esercizi reali, soluzioni ragionate, errori tipici e criteri di autovalutazione. Quando una sezione promette quiz, test o simulazioni, deve contenere le domande effettive e non istruzioni generiche su come studiare. Mantieni separati quesiti e soluzioni, verifica il numero richiesto, evita duplicati e non inventare regole d'esame non verificate.",
        "Narrativo": "Sviluppa personaggi, conflitto, cause e conseguenze in scene concrete. Ogni capitolo deve avere una funzione narrativa distinta.",
        "Romanzo Classico": "Usa una costruzione narrativa solida, personaggi coerenti, ambientazione e temi sviluppati attraverso azioni e dialoghi; evita imitazioni di autori viventi.",
        "Contemporaneo": "Racconta conflitti e relazioni con voce naturale, dettagli specifici e temi attuali trattati attraverso la storia, non con prediche.",
        "Self-Help": "Definisci problemi realistici, pratiche graduali, esempi e criteri di verifica. Evita promesse di trasformazione garantita o consigli clinici.",
        "Manuale Pratico": "Fornisci un percorso eseguibile: materiali o prerequisiti, passaggi, controlli, errori, alternative e risultato finale verificabile.",
        "Storico": "Ordina il racconto per nessi causali e cronologia, distinguendo fonti, fatti, interpretazioni e controversie. Non inventare citazioni o date."
    }
    return profili.get(genere, "Mantieni una struttura coerente con pubblico, obiettivo, genere e limiti dichiarati.")


def estrai_numero_ricette(titolo, trama, obiettivo):
    testo = f"{titolo} {trama} {obiettivo}".lower()
    match = re.search(r"\\b(\\d{1,3})\\s+(?:ricette|recipes|recetas|recettes|rezepte)\\b", testo)
    return int(match.group(1)) if match else None


def profilo_struttura_indice(genere, titolo, trama, obiettivo):
    """Evita che una stessa gabbia 15-18 capitoli venga applicata a libri incompatibili."""
    if genere == "Ricettario":
        numero = estrai_numero_ricette(titolo, trama, obiettivo)
        quantità = f"esattamente {numero}" if numero else "un numero coerente con la richiesta"
        return f"""RICETTARIO: crea {quantità} ricette effettive, distribuite in parti tematiche coerenti. Ogni ricetta è un Capitolo autonomo e completo. Se è richiesto un numero preciso di ricette, crea esattamente quel numero di Capitoli e ciascun Capitolo deve avere il nome di una ricetta: non usare Capitoli per introduzione, ingredienti, attrezzatura, tecniche o consigli. Le Parti possono orientare il lettore senza aggiungere Capitoli introduttivi. Non creare sottocapitoli 1.1, 1.2 o 1.3 per espandere la stessa ricetta. Il numero delle ricette nell'indice deve coincidere con il numero richiesto."""
    if genere in {"Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Narrativo", "Romanzo Classico", "Contemporaneo", "Biografia"}:
        return "NARRATIVA E BIOGRAFIA: organizza 3-6 Parti e un numero di capitoli proporzionato all'arco narrativo. Non imporre sottocapitoli a ogni capitolo: usali solo se sono necessari e non spezzano artificialmente scene o svolte. Ogni titolo deve nominare una scena, una scelta, un luogo, un personaggio, un oggetto o una conseguenza specifici del brief. Almeno un terzo dei titoli deve contenere parole concrete tratte dal titolo o dalla trama. Evita titoli generici come 'Il ritorno', 'La scoperta', 'L'incontro inaspettato', 'Il richiamo del passato', 'Riflessioni' o 'La fine'."
    if genere in {"Quiz Scientifico", "Test Prep (Preparazione Esami)"}:
        return "QUIZ E TEST PREP: organizza fondamenti, esercitazione graduata, simulazioni e correzioni. Ogni unità deve indicare una competenza verificabile; non creare capitoli riempitivi."
    return "SAGGISTICA E MANUALI: usa 4-6 Parti, 15-18 capitoli effettivi e 3-5 sottocapitoli solo quando corrispondono a concetti o passaggi realmente distinti."


def normalizza_indice_generato(indice):
    """Rimuove solo rumore di formattazione, senza alterare l'architettura proposta."""
    righe = []
    for riga in (indice or "").splitlines():
        pulita = re.sub(r"^\s*[-*#]+\s*", "", riga).strip()
        if pulita.lower() in {"indice", "table of contents", "sommaire", "inhaltsverzeichnis"}:
            continue
        if pulita:
            righe.append(pulita)
    return "\n".join(righe).strip()


def criticita_indice_generato(indice, genere, titolo, trama, obiettivo):
    """Controllo deterministico leggero: intercetta gli errori che il modello tende a ripetere."""
    testo = normalizza_indice_generato(indice)
    righe = testo.splitlines()
    capitoli = [riga for riga in righe if re.match(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+", riga)]
    parti = [riga for riga in righe if re.match(r"(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\s+", riga)]
    if not capitoli:
        return ["non sono stati riconosciuti capitoli nel formato richiesto"]

    problemi = []
    narrativi = {"Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Narrativo", "Romanzo Classico", "Contemporaneo", "Biografia"}
    if genere != "Ricettario" and len(parti) < 4:
        problemi.append(f"struttura troppo breve: sono presenti solo {len(parti)} Parti")
    minimo_capitoli = 12 if genere not in {"Ricettario"} else 0
    if len(capitoli) < minimo_capitoli:
        problemi.append(f"struttura troppo breve: sono presenti solo {len(capitoli)} Capitoli, ne servono almeno {minimo_capitoli}")
    if genere not in narrativi and genere != "Ricettario":
        capitoli_senza_sviluppo = []
        for posizione, capitolo in enumerate(capitoli):
            inizio = righe.index(capitolo)
            fine = next((i for i in range(inizio + 1, len(righe)) if re.match(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+", righe[i]) or re.match(r"(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\s+", righe[i])), len(righe))
            sottosezioni = sum(1 for riga in righe[inizio + 1:fine] if re.match(r"^\d+\.\d+\s+", riga))
            if sottosezioni < 2:
                capitoli_senza_sviluppo.append(capitolo)
        if capitoli_senza_sviluppo:
            problemi.append("capitoli senza almeno due sottocapitoli distinti: " + "; ".join(capitoli_senza_sviluppo[:3]))

    if genere == "Ricettario":
        richieste = estrai_numero_ricette(titolo, trama, obiettivo)
        if richieste and len(capitoli) != richieste:
            problemi.append(f"sono richieste {richieste} ricette, ma l'indice contiene {len(capitoli)} capitoli")
        titoli_capitoli = " ".join(capitoli).lower()
        non_ricette = ("introduzione", "ingredient", "attrezz", "tecniche", "consigli", "dispensa")
        if any(parola in titoli_capitoli for parola in non_ricette):
            problemi.append("un capitolo del ricettario è introduttivo o tecnico invece di essere una ricetta")

    if genere in narrativi:
        titoli_generici = {
            "il ritorno", "la scoperta", "l'inizio", "la fine", "il conflitto", "la scelta", "la crisi",
            "riflessioni", "sogni e memorie", "nuovi inizi", "l'incontro inaspettato", "il richiamo del passato",
            "la dolcezza del ricordo", "il richiamo della tradizione", "riscoprire se stessi", "la verità", "il segreto"
        }
        trovati = []
        for capitolo in capitoli:
            nome = re.sub(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+\s*:\s*", "", capitolo).strip().lower()
            if nome in titoli_generici:
                trovati.append(capitolo)
        if len(trovati) >= 2:
            problemi.append("titoli narrativi troppo generici: " + "; ".join(trovati[:3]))
        parole_da_escludere = {
            "della", "delle", "dello", "degli", "dalla", "nelle", "nello", "come", "con", "una", "uno", "per",
            "che", "del", "dei", "gli", "le", "il", "la", "un", "e", "di", "da", "in", "su", "tra", "fra",
            "storia", "romanzo", "guida", "raccontare", "lettore", "lettori", "obiettivo", "titolo", "libro"
        }
        parole_brief = {
            parola for parola in re.findall(r"[a-zàèéìòóù]{4,}", f"{titolo} {trama}".lower())
            if parola not in parole_da_escludere
        }
        titoli_con_ancora = 0
        for capitolo in capitoli:
            nome = re.sub(r"(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+\s*:\s*", "", capitolo).lower()
            if any(parola in nome for parola in parole_brief):
                titoli_con_ancora += 1
        soglia = max(3, (len(capitoli) + 2) // 3)
        if titoli_con_ancora < soglia:
            problemi.append(
                f"titoli narrativi poco ancorati agli elementi concreti del brief ({titoli_con_ancora}/{len(capitoli)} titoli specifici)"
            )
    if genere in {"Quiz Scientifico", "Test Prep (Preparazione Esami)"}:
        testo_minuscolo = testo.lower()
        if "quiz" not in testo_minuscolo and "domand" not in testo_minuscolo:
            problemi.append("manca una sezione con quiz o domande effettive")
        if "simulaz" not in testo_minuscolo:
            problemi.append("manca una sezione di simulazione")
    return problemi


def audit_editoriale_indice_generato(indice, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov):
    """Usa esattamente lo stesso metro del pulsante 'Voto indice', evitando approvazioni incoerenti."""
    risposta = valuta_indice_editoriale(
        indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, ""
    ).strip()
    match = re.search(r"(?im)^\s*(?:voto\s+complessivo|voto)\s*:\s*(10|[0-9])\s*(?:/\s*10)?\b", risposta)
    voto = int(match.group(1)) if match else 0
    difetti = re.search(r"(?ims)^\s*(?:miglioramenti consigliati|difetti)\s*:\s*(.+?)(?:\n\s*[A-ZÀ-Ú][A-ZÀ-Ú ]+\s*:|$)", risposta)
    return voto, (difetti.group(1).strip() if difetti else risposta)


def genera_indice_controllato(prompt, system_prompt, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov):
    """Genera l'indice e applica fino a due revisioni su vincoli oggettivi e qualità editoriale."""
    corrente = normalizza_indice_generato(chiedi_gpt(prompt, system_prompt))
    for tentativo in range(3):
        problemi = criticita_indice_generato(corrente, genere, titolo, trama, obiettivo)
        voto_editoriale, difetti_editoriali = (0, "")
        if not problemi:
            voto_editoriale, difetti_editoriali = audit_editoriale_indice_generato(
                corrente, genere, titolo, trama, obiettivo, lingua, stile, narrativa, pov
            )
            if voto_editoriale == 10:
                esito = "Indice approvato: 10/10 nel controllo strutturale ed editoriale automatico."
                if tentativo:
                    esito = f"Indice corretto automaticamente al controllo {tentativo} e approvato 10/10."
                st.session_state["ultimo_controllo_indice"] = esito
                return corrente
            problemi.append(f"audit editoriale {voto_editoriale}/10: {difetti_editoriali}")
        if tentativo == 2:
            st.session_state["ultimo_controllo_indice"] = "Attenzione: l'indice non ha raggiunto 10/10 e richiede una verifica manuale: " + "; ".join(problemi)
            return ""
        revisione = prompt + f"""

REVISIONE OBBLIGATORIA DELL'INDICE — TENTATIVO {tentativo + 1}
La bozza precedente non rispetta questi vincoli oggettivi/editoriali: {'; '.join(problemi)}.
Riscrivi l'intero indice, senza commenti e senza la parola 'Indice' in apertura. Correggi tutti i punti segnalati;
non limitarti a rinominare i titoli. Mantieni soltanto argomenti attinenti al brief.
"""
        corrente = normalizza_indice_generato(chiedi_gpt(revisione, system_prompt))
    return ""


def tipo_sezione_editoriale(sezione):
    pulita = sezione.strip()
    if re.match(r'(?i)^(parte|part|partie|teil|partea|часть|الجزء|部分)\b', pulita):
        return "parte"
    if re.match(r'(?i)^(capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+', pulita):
        return "capitolo"
    if re.match(r'^\d+\.\d+\s+', pulita):
        return "sottocapitolo"
    return "frontespizio"


def chiave_sezione(sezione):
    return f"txt_{sezione.replace(' ', '_').replace('.', '')}"


def sezioni_mancanti_per_esportazione(sezioni, genere):
    """Non consente di esportare un libro se l'indice contiene sezioni non effettivamente redatte."""
    mancanti = []
    minimi = {"parte": 35, "capitolo": 90 if genere == "Ricettario" else 120, "sottocapitolo": 120, "frontespizio": 40}
    for sezione in sezioni:
        testo = pulisci_testo_editoriale(st.session_state.get(chiave_sezione(sezione), "")).strip()
        if len(testo.split()) < minimi[tipo_sezione_editoriale(sezione)]:
            mancanti.append(sezione)
    return mancanti


def genera_sezione_con_ripetizione(prompt, system_prompt, sezione, lingua, tentativi=2):
    """Riprova una sezione senza perdere le precedenti; evita libri interrotti a metà dopo un errore transitorio."""
    ultimo_errore = None
    for tentativo in range(1, tentativi + 1):
        try:
            testo = chiedi_gpt(prompt, system_prompt)
            if not testo or testo.startswith("ERRORE:"):
                raise RuntimeError(testo or "Risposta vuota")
            # La stesura non effettua automaticamente una ricerca: la verifica
            # viene attivata solo dal motore editoriale quando individua fatti variabili.
            return pulisci_testo_editoriale(testo)
        except Exception as exc:
            ultimo_errore = exc
    raise RuntimeError(f"Impossibile completare la sezione dopo {tentativi} tentativi: {ultimo_errore}")


def capitolo_padre(indice, sezione):
    """Restituisce il capitolo che contiene un sottocapitolo, utile per riconoscere le simulazioni."""
    padre = ""
    confronto = re.sub(r"\s+", " ", sezione.strip().lower())
    for riga in (indice or "").splitlines():
        pulita = riga.strip()
        if re.match(r"(?i)^capitolo\s+\d+", pulita):
            padre = pulita
        if re.sub(r"\s+", " ", pulita.lower()) == confronto:
            return padre
    return padre


def numero_domande_simulazione(indice, trama, obiettivo):
    """Ricava un conteggio soltanto quando il brief lo dichiara; evita numeri inventati."""
    testo = f"{indice}\n{trama}\n{obiettivo}"
    corrispondenze = re.findall(
        r"(?is)(?:simulazion\w*|test\s+completo|prova\s+completa).{0,90}?(\d{1,3})\s+(?:domand\w*|quesit\w*)"
        r"|(\d{1,3})\s+(?:domand\w*|quesit\w*).{0,90}?(?:simulazion\w*|test\s+completo|prova\s+completa)",
        testo
    )
    numeri = [int(a or b) for a, b in corrispondenze if (a or b)]
    return max(numeri) if numeri else 0


def sezione_simulazione_test_prep(sezione, indice, genere):
    """Attiva il generatore a blocchi solo nella sezione che deve contenere i quesiti della simulazione."""
    if genere != "Test Prep (Preparazione Esami)":
        return False
    titolo = sezione.lower()
    padre = capitolo_padre(indice, sezione).lower()
    è_simulazione = "simulazione" in titolo or "simulazione" in padre
    è_contenuto_test = any(parola in titolo for parola in ("domande", "quesiti", "quiz", "esecuzione", "test"))
    return è_simulazione and è_contenuto_test


def conta_domande_test_prep(testo):
    return len(re.findall(r"(?im)^\s*domanda\s+\d{1,3}\s*[:.-]", testo or ""))


def domande_normalizzate_test_prep(testo):
    domande = re.findall(r"(?im)^\s*domanda\s+\d{1,3}\s*[:.-]\s*(.+)$", testo or "")
    return [re.sub(r"[^a-z0-9àèéìòóù ]", "", domanda.lower()).strip() for domanda in domande]


def genera_simulazione_test_prep(prompt_base, system_prompt, sezione, indice, trama, obiettivo, lingua):
    """Genera prove lunghe in blocchi verificabili, evitando simulazioni promesse ma incomplete."""
    totale = numero_domande_simulazione(indice, trama, obiettivo)
    if totale < 10:
        return genera_sezione_con_ripetizione(prompt_base, system_prompt, sezione, lingua)

    dimensione_blocco = 10
    blocchi_domande, domande_precedenti = [], []
    for inizio in range(1, totale + 1, dimensione_blocco):
        fine = min(inizio + dimensione_blocco - 1, totale)
        vincolo = f"""
Questa è la parte domande della simulazione '{sezione}', gruppo {inizio}-{fine} di {totale}.
Genera ESATTAMENTE {fine - inizio + 1} quesiti originali, numerati da DOMANDA {inizio:02d}: a DOMANDA {fine:02d}:.
Per ogni domanda usa quattro opzioni A), B), C), D). NON scrivere risposte, soluzioni, commenti,
punteggi o istruzioni su come prepararsi. Distribuisci i quesiti sui contenuti obbligatori dell'indice
e del brief; non ripetere le domande già prodotte qui sotto.
Domande già prodotte: {' | '.join(domande_precedenti) or 'nessuna'}
"""
        blocco, nuove_domande = "", []
        for tentativo_blocco in range(3):
            correzione = ""
            if tentativo_blocco:
                correzione = (
                    "\nCORREZIONE OBBLIGATORIA: la bozza precedente era incompleta o ripeteva quesiti già usati. "
                    "Sostituiscila interamente con domande nuove, numerate nel formato richiesto e senza testo introduttivo."
                )
            candidato = genera_sezione_con_ripetizione(prompt_base + vincolo + correzione, system_prompt, sezione, lingua)
            domande_candidate = domande_normalizzate_test_prep(candidato)
            conteggio_corretto = conta_domande_test_prep(candidato) == (fine - inizio + 1)
            senza_duplicati = not (set(domande_candidate) & set(domande_precedenti))
            if conteggio_corretto and senza_duplicati:
                blocco, nuove_domande = candidato, domande_candidate
                break
            blocco, nuove_domande = candidato, domande_candidate
        if conta_domande_test_prep(blocco) != (fine - inizio + 1) or (set(nuove_domande) & set(domande_precedenti)):
            st.session_state["avviso_simulazione_test_prep"] = (
                f"La simulazione '{sezione}' richiede una verifica: il blocco {inizio}-{fine} non ha superato "
                "il controllo automatico di quantità o unicità. Il controllo coerenza indicherà le correzioni necessarie."
            )
        blocchi_domande.append(blocco)
        domande_precedenti.extend(nuove_domande)

    corpo_domande = "\n\n".join(blocchi_domande)
    prompt_chiave = f"""Crea la chiave delle soluzioni per la simulazione '{sezione}' in lingua {lingua}.
Le domande seguenti sono già state redatte. Fornisci ESATTAMENTE una riga di soluzione per ogni
DOMANDA da 01 a {totale:02d}, con questo formato: SOLUZIONE 01: lettera corretta - spiegazione breve e concreta.
Non riscrivere le domande, non aggiungere nuove domande, non usare link o fonti e non omettere numeri.

DOMANDE DELLA SIMULAZIONE
{corpo_domande}
"""
    chiave = ""
    for tentativo_chiave in range(3):
        correzione = "" if not tentativo_chiave else "\nCORREZIONE OBBLIGATORIA: inserisci tutte e sole le soluzioni numerate richieste."
        chiave = genera_sezione_con_ripetizione(prompt_chiave + correzione, system_prompt, sezione, lingua)
        if len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", chiave or "")) == totale:
            break
    if len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", chiave or "")) != totale:
        st.session_state["avviso_simulazione_test_prep"] = (
            f"La simulazione '{sezione}' richiede una verifica: la chiave delle soluzioni non ha il conteggio previsto. "
            "Il controllo coerenza indicherà le correzioni necessarie."
        )
    return (
        f"SIMULAZIONE: DOMANDE\n\n{corpo_domande}\n\n"
        f"SOLUZIONI COMMENTATE - CONSULTALE SOLO DOPO AVER COMPLETATO LA PROVA\n\n{chiave}"
    )


def criticita_specificita(testo, genere, sezione):
    """Individua bozze genericamente motivazionali prima che finiscano nel manoscritto."""
    pulito = pulisci_testo_editoriale(testo or "").strip()
    parole = pulito.split()
    if tipo_sezione_editoriale(sezione) == "parte":
        return ""
    if len(parole) < 150:
        return "testo troppo breve per sviluppare l'argomento assegnato"

    basso = pulito.lower()
    formule_generiche = (
        "è fondamentale", "e fondamentale", "è cruciale", "e cruciale", "in modo efficace",
        "è importante", "e importante", "con sicurezza", "molto utile", "potente strumento"
    )
    genericita = sum(basso.count(formula) for formula in formule_generiche)
    segnali_per_genere = {
        "Saggio Scientifico": ("definiz", "evidenz", "limite", "esempio"),
        "Quiz Scientifico": ("domanda", "risposta", "spiegazione", "errore"),
        "Manuale Tecnico": ("passo", "verifica", "errore", "esempio"),
        "Religioso / Teologico": ("testo", "tradizion", "interpret", "contesto"),
        "Spirituale / Esoterico": ("pratica", "esperienza", "limite", "esercizio"),
        "Meditazione / Mindfulness": ("esercizio", "respiro", "osserv", "durata"),
        "Business & Marketing": ("caso", "metrica", "azione", "cliente"),
        "Economia e Finanza": ("dato", "rischio", "scenario", "esempio"),
        "Romanzo Rosa": ("dialog", "scena", "personagg", "relazione"),
        "Thriller / Noir": ("scena", "indizio", "conflitto", "personagg"),
        "Fantasy": ("scena", "personagg", "conflitto", "mondo"),
        "Fantascienza": ("scena", "personagg", "conseguenz", "tecnolog"),
        "Manuale Psicologico": ("esercizio", "esempio", "limite", "pratica"),
        "Biografia": ("evento", "contesto", "periodo", "scelta"),
        "Ricettario": ("ingredient", "procedimento", "cottura", "porzion"),
        "Test Prep (Preparazione Esami)": ("domanda", "risposta", "esercizio", "errore"),
        "Narrativo": ("scena", "personagg", "azione", "dialog"),
        "Romanzo Classico": ("scena", "personagg", "azione", "dialog"),
        "Contemporaneo": ("scena", "personagg", "azione", "dialog"),
        "Self-Help": ("esercizio", "passo", "esempio", "verifica"),
        "Manuale Pratico": ("passo", "material", "errore", "risultato"),
        "Storico": ("evento", "contesto", "fonte", "periodo")
    }
    segnali = segnali_per_genere.get(genere, ())
    trovati = sum(1 for segnale in segnali if segnale in basso)
    if genericita >= 5 and trovati < 2:
        return "eccesso di formule generiche senza esempi, scene, dati, procedure o strumenti specifici del genere"
    return ""


def genera_contenuto_editoriale(prompt, system_prompt, sezione, indice, trama, genere, obiettivo, lingua):
    """Mantiene il flusso comune per tutti i generi e applica la logica speciale solo quando serve."""
    if sezione_simulazione_test_prep(sezione, indice, genere):
        return genera_simulazione_test_prep(prompt, system_prompt, sezione, indice, trama, obiettivo, lingua)
    testo = genera_sezione_con_ripetizione(prompt, system_prompt, sezione, lingua)
    criticita = criticita_specificita(testo, genere, sezione)
    if criticita:
        testo = genera_sezione_con_ripetizione(
            prompt + f"""

REVISIONE OBBLIGATORIA DI QUALITÀ
La prima bozza è stata rifiutata perché presenta: {criticita}.
Riscrivi integralmente la sezione. Ogni paragrafo deve aggiungere un fatto, una scena, una procedura,
un esempio, un caso, un esercizio, un dato o una conseguenza specifica del genere '{genere}'.
Elimina frasi motivazionali, definizioni vaghe e ripetizioni. Non descrivere ciò che il lettore potrebbe fare:
mostra il contenuto concreto richiesto dal titolo della sezione.
""",
            system_prompt, sezione, lingua
        )
    # Le ricerche web sono riservate a leggi, prezzi, versioni, requisiti e altri
    # dati soggetti a cambiamento; i contenuti didattici stabili non consumano credito web.
    if richiede_verifica_fatti(testo, sezione):
        testo = verifica_e_correggi_fatti_online(testo, sezione, lingua)
    return testo

# NUOVA FUNZIONE: Motore Decisionale per attivare i 3 Cervelli in base alla Sidebar
def valuta_approccio_neurologico(genere, stile, narrativa):
    """
    Decide se l'argomento e lo stile richiedono la manipolazione dei 3 cervelli
    o un approccio più analitico/oggettivo.
    """
    trigger_neuro_stile = ["Persuasivo (Neuromarketing Applicato)", "Conversazionale ed Empatico", "Storytelling Immersivo", "Epico ed Evocativo"]
    trigger_neuro_narrativa = ["Coinvolgente e Narrativo", "Ispirazionale e Motivante", "Storytelling Emozionale", "Diretto e Pratico (Action-oriented)"]
    trigger_neuro_genere = ["Business & Marketing", "Economia e Finanza", "Manuale Psicologico", "Romanzo Rosa", "Thriller / Noir", "Spirituale / Esoterico"]
    
    if stile in trigger_neuro_stile or narrativa in trigger_neuro_narrativa or genere in trigger_neuro_genere:
        return True
    return False

# ======================================================================================================================
# 6. SIDEBAR: SETUP EDITORIALE AVANZATO E CARICAMENTO FONTI
# ======================================================================================================================
with st.sidebar:
    lingua_sel = st.selectbox("🌐 Lingua / Language", list(TRADUZIONI.keys()))
    L = TRADUZIONI.get(lingua_sel, TRADUZIONI["Italiano"])
    st.title(L["side_tit"])
    val_titolo = st.text_input(L["lbl_tit"])
    val_autore = st.text_input(L["lbl_auth"])
    
    # --- NUOVA SEZIONE CARICAMENTO FONTI ---
    st.markdown("### 📂 Fonti Esterne (Opzionale)")
    st.markdown("<small>Carica PDF o DOCX per aiutare l'IA nel ragionamento di stesura.</small>", unsafe_allow_html=True)
    file_caricati = st.file_uploader("Carica Fonti Esterne", type=['pdf', 'docx'], accept_multiple_files=True, label_visibility="collapsed")
    if file_caricati:
        if len(file_caricati) > 10:
            st.warning("Hai superato il limite di 10 file. Verranno analizzati i primi 10.")
            file_caricati = file_caricati[:10]
        with st.spinner("Lettura e analisi fonti in corso..."):
            st.session_state["conoscenza_extra"] = estrai_testo_da_files(file_caricati)
            if st.session_state["conoscenza_extra"]:
                st.success(f"Analizzati {len(file_caricati)} documenti. Pronti per l'uso!")
    else:
        st.session_state["conoscenza_extra"] = ""
    
    st.markdown("---")
    # --- AGGIUNTA "STORICO" AI GENERI ---
    lista_gen = ["Saggio Scientifico", "Quiz Scientifico", "Manuale Tecnico", "Religioso / Teologico", "Spirituale / Esoterico", "Meditazione / Mindfulness", "Business & Marketing", "Economia e Finanza", "Romanzo Rosa", "Thriller / Noir", "Fantasy", "Fantascienza", "Manuale Psicologico", "Biografia", "Ricettario", "Test Prep (Preparazione Esami)", "Narrativo", "Romanzo Classico", "Contemporaneo", "Self-Help", "Manuale Pratico", "Storico"]
    val_genere = st.selectbox(L["lbl_gen"], lista_gen)
    
    stili_estesi = [
        "Standard", 
        "Professionale Accademico", 
        "Persuasivo (Neuromarketing Applicato)", 
        "Conversazionale ed Empatico", 
        "Scientifico Divulgativo", 
        "Storytelling Immersivo", 
        "Giornalistico d'Inchiesta", 
        "Socratico (Dialogico / Riflessivo)", 
        "Epico ed Evocativo", 
        "Minimalista ed Essenziale"
    ]
    val_stile = st.selectbox(L["lbl_style"], stili_estesi)

    direttive_indice_tipologia = {
        "Standard": "Crea un percorso lineare in 5-6 Parti e 15-18 Capitoli: basi, sviluppo progressivo, applicazione, verifica e sintesi. Ogni sottocapitolo deve avere un obiettivo concreto e un risultato leggibile.",
        "Professionale Accademico": "Organizza l'indice in contesto, definizioni, quadro teorico, metodologia, analisi, evidenze, limiti, implicazioni e riferimenti. Separa chiaramente ipotesi, dati, metodo e risultati; prevedi criteri di valutazione e fonti.",
        "Persuasivo (Neuromarketing Applicato)": "Costruisci il percorso da problema e consapevolezza a soluzione, prove, obiezioni, benefici, applicazione e azione. Inserisci casi, comparazioni e piani d'azione senza promesse garantite o claim non verificabili.",
        "Conversazionale ed Empatico": "Sequenzia l'indice come un accompagnamento: situazione del lettore, ostacoli, spiegazione semplice, esercitazione guidata, verifica e autonomia. Usa domande guida, riepiloghi e passaggi graduali senza infantilizzare.",
        "Scientifico Divulgativo": "Procedi da basi e contesto a meccanismi, evidenze, esempi, applicazioni e limiti. Indica dove servono fonti aggiornate, distingui fatti, ipotesi e analogie e inserisci esperimenti mentali o verifiche pratiche quando pertinenti.",
        "Storytelling Immersivo": "Progetta un arco narrativo completo con situazione iniziale, personaggi, desiderio, conflitto, ostacoli, svolte, conseguenze, climax e risoluzione. Ogni capitolo deve modificare la situazione o approfondire un personaggio, evitando capitoli di riempimento.",
        "Giornalistico d'Inchiesta": "Organizza il percorso da domanda iniziale a contesto, fonti primarie, testimonianze, verifiche indipendenti, contraddizioni, prove, responsabilità e conclusioni. Specifica quali fatti devono essere documentati e separa dati accertati da ipotesi.",
        "Socratico (Dialogico / Riflessivo)": "Costruisci l'indice attraverso domande progressive: presupposti, definizioni, dubbi, obiezioni, esempi, conseguenze e sintesi. Ogni capitolo deve porre una domanda centrale e chiuderla con una risposta argomentata o un esercizio di riflessione.",
        "Epico ed Evocativo": "Crea una progressione ampia con origine, chiamata, prove, alleati, opposizioni, trasformazione, crisi, compimento e significato finale. Mantieni immagini evocative nei titoli ma indica sempre un contenuto concreto e coerente con il genere.",
        "Minimalista ed Essenziale": "Riduci il libro a 15-18 capitoli indispensabili, con un solo obiettivo per capitolo e sottocapitoli non sovrapposti. Usa titoli brevi e operativi, elimina digressioni e assegna a ogni sezione un risultato verificabile."
    }
    direttiva_indice_selezionata = direttive_indice_tipologia.get(val_stile, direttive_indice_tipologia["Standard"])
    
    st.markdown("---")
    # --- AGGIUNTA "STORICO E DOCUMENTALE" AGLI STILI DI RACCONTO ---
    val_narrativa = st.selectbox(L["lbl_narrative"], [
        "Coinvolgente e Narrativo", "Tecnico e Analitico", "Ispirazionale e Motivante", 
        "Socratico (Domanda/Risposta)", "Storytelling Emozionale", "Diretto e Pratico (Action-oriented)", "Storico e Documentale"
    ])
    
    # NUOVO BLOCCO: Punto di Vista (POV)
    lista_pov = [
        "Tu (Diretto, confidenziale e personale)",
        "Voi (Plurale, autorevole e rispettoso)",
        "Noi (Inclusivo, partecipativo e didattico)",
        "Impersonale / Terza Persona (Distaccato, analitico, oggettivo)"
    ]
    val_pov = st.selectbox(L.get("lbl_pov", "Punto di Vista (Pronome)"), lista_pov)
    
    # Definizioni disponibili prima del loro primo utilizzo nella UI.
    # Restano presenti anche nel modulo di memoria sottostante per compatibilità.
    def costruisci_specifica_editoriale(titolo, genere, stile, narrativa, pov, obiettivo, argomento, approfondimenti=""):
        return f"""=== SPECIFICA EDITORIALE STRUTTURATA ===
Titolo: {titolo}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}

OBIETTIVO OPERATIVO:
{obiettivo}

ARGOMENTO E CONFINI:
{argomento}

APPROFONDIMENTI PRIORITARI (FACOLTATIVI):
{approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

Per ogni sezione ricava un risultato concreto, il livello del lettore, i concetti necessari,
gli esempi o le procedure da produrre e ciò che deve restare fuori per evitare ripetizioni.
"""

    def analizza_coerenza_libro(indice, contenuti, obiettivo, argomento):
        risultati = ["REPORT CONTROLLO COERENZA DEL LIBRO"]
        testo = "\n".join(contenuti.values()) if contenuti else ""
        capitoli = re.findall(r"(?im)^(?:Capitolo|Chapter|CAPITOLO)\\s+\\d+", indice or "")
        sottocapitoli = re.findall(r"(?m)^\\d+\\.\\d+\\s+", indice or "")
        risultati.append(f"Capitoli rilevati: {len(capitoli)}")
        risultati.append(f"Sottocapitoli rilevati: {len(sottocapitoli)}")
        if not indice.strip(): risultati.append("ERRORE: indice assente")
        if not obiettivo.strip(): risultati.append("AVVISO: obiettivo assente")
        if not argomento.strip(): risultati.append("AVVISO: argomento assente")
        if len(testo.strip()) < 1000: risultati.append("AVVISO: contenuto ancora troppo breve per una verifica completa")
        frasi = [f.strip().lower() for f in re.split(r"[.!?]+", testo) if len(f.strip()) > 40]
        duplicati = len(frasi) - len(set(frasi))
        risultati.append(f"Frasi duplicate identiche rilevate: {max(0, duplicati)}")
        if duplicati == 0: risultati.append("OK: nessuna duplicazione identica rilevata nel testo disponibile")
        return "\n".join(risultati)

    val_goal = st.text_input(L["lbl_goal"], placeholder="Es: Mantenere l'attenzione alta, far emozionare...")
    val_trama = st.text_area(L["lbl_plot"], height=150)
    val_approfondimenti = st.text_area(
        "Approfondimenti (facoltativo)",
        height=130,
        placeholder="Inserisci istruzioni, aspetti da trattare con maggiore attenzione, vincoli, esempi o temi obbligatori."
    )
    specifica_editoriale = costruisci_specifica_editoriale(
        val_titolo, val_genere, val_stile, val_narrativa, val_pov, val_goal, val_trama, val_approfondimenti
    )
    
    # PULSANTE RESET BLINDATO: Unico modo per svuotare la session_state
    if st.button(L["btn_res"]):
        for key in list(st.session_state.keys()): del st.session_state[key]
        st.rerun()

# ======================================================================================================================
# 7. LOGICA DI MEMORIA E COERENZA (EVITA RIPETIZIONI GLOBALI) E INTEGRAZIONE FONTI
# ======================================================================================================================
def genera_contesto_avanzato(sezione_corrente):
    contesto = ""
    if st.session_state.get("conoscenza_extra"):
        contesto += f"=== FONTI ESTERNE DI RIFERIMENTO (USATE PER RAGIONAMENTO) ===\n{st.session_state['conoscenza_extra'][:8000]}\n\n"
        
    for s in st.session_state.get("lista_capitoli", []):
        if s == sezione_corrente: break
        k = f"txt_{s.replace(' ', '_').replace('.', '')}"
        if k in st.session_state and st.session_state[k].strip():
            # Memoria estesa: il riepilogo breve da 150 caratteri non era sufficiente
            # per distinguere concetti, esempi e procedure già utilizzati.
            testo_precedente = st.session_state[k]
            contesto += f"- Trattato in {s}:\n{testo_precedente[:1200]}\n"
    return contesto

def individua_sottocapitoli_del_capitolo(capitolo, sezioni):
    """Restituisce soltanto i sottocapitoli numerati appartenenti al capitolo selezionato."""
    match = re.match(r"(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+(\d+)", capitolo.strip())
    if not match:
        return []
    numero = match.group(1)
    return [s for s in sezioni if re.match(rf"^{re.escape(numero)}\.\d+\s", s.strip())]

def individua_sezioni_da_stendere(sezioni):
    """Restituisce capitoli e sottocapitoli in ordine, escludendo Parti e appendici."""
    regex_capitolo = r'(?i)^(?:capitolo|chapter|kapitel|capítulo|chapitre|capitolul|глава|الفصل|章节)\s+\d+'
    regex_sottocapitolo = r'^\d+\.\d+\s+'
    return [
        sezione for sezione in sezioni
        if re.match(regex_capitolo, sezione.strip()) or re.match(regex_sottocapitolo, sezione.strip())
    ]

def crea_prompt_stesura_sezione(sezione, indice, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti=""):
    """Costruisce il prompt comune usato dalla stesura singola e dalla stesura di un capitolo intero."""
    memoria = genera_contesto_avanzato(sezione)
    tipo_sezione = tipo_sezione_editoriale(sezione)
    profilo_genere = profilo_genere_stesura(genere)
    profilo_tipologia = profilo_tipologia_stesura(stile)
    regola_struttura = profilo_struttura_indice(genere, "", trama, obiettivo)
    direttiva_test_prep = ""
    if genere == "Test Prep (Preparazione Esami)":
        direttiva_test_prep = """
- Se il titolo della sezione contiene quiz, test, domande, autovalutazione o esercizi, crea il materiale promesso: quesiti originali con quattro opzioni, risposta corretta e spiegazione della scelta. Non limitarti a spiegare come affrontare un quiz.
- Se il titolo contiene simulazione, genera solo ciò che il titolo assegna alla sezione; la sezione che contiene domande o esecuzione sarà completata dal generatore in blocchi con quesiti e soluzioni separati.
- Evita formule generiche sullo studio: ogni paragrafo deve introdurre una competenza d'esame, un errore concreto, una procedura, un quesito o una decisione verificabile.
"""
    return f"""
INDICE GENERALE (STUDIALO PER CAPIRE COSA NON DEVI ANTICIPARE):
{indice}

MEMORIA CONTENUTI PRECEDENTI (Per non ripetersi):
{memoria}

=== PARAMETRI EDITORIALI SARTORIALI (DA APPLICARE TASSATIVAMENTE IN QUESTA SEZIONE) ===
- Argomento Centrale / Trama: {trama}
- Genere Letterario: {genere}
- Tipologia di Scrittura: {stile}
- Stile di Racconto: {narrativa}
- Punto di Vista (POV): {pov}
- Obiettivo Emozionale/Pratico: {obiettivo}
- Approfondimenti prioritari: {approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

=== PROFILO EDITORIALE DA RISPETTARE ===
- Regole del genere: {profilo_genere}
- Regole della tipologia di scrittura: {profilo_tipologia}
- Regole della struttura: {regola_struttura}
- Tipo della sezione corrente: {tipo_sezione}

AZIONE:
Scrivi ora la sezione ESATTA: '{sezione}'. Il testo deve essere rigorosamente in lingua {lingua}.
- Se la sezione è una Parte, scrivi soltanto una breve apertura che spiega lo scopo della Parte e come usarla: non sviluppare o riassumere i capitoli che seguono.
- Se la sezione è un Capitolo con sottocapitoli nell'indice, non anticipare né risolvere gli argomenti assegnati ai relativi sottocapitoli.
- Se il genere è Ricettario, ogni Capitolo che porta il nome di un piatto è una sola ricetta completa: non aggiungere sottocapitoli autonomi e non ripetere una ricetta già presente nella Parte o in altre sezioni.
{direttiva_test_prep}
- Rispetta integralmente i parametri editoriali e usa tassativamente il POV richiesto ({pov}).
- Tratta con priorità gli approfondimenti forniti, ma soltanto nelle sezioni cui sono pertinenti; non ripeterli artificialmente e non anticipare contenuti assegnati a sezioni successive.
- Sii profondo ed esaustivo nell'ambito della sezione, senza rubare materiale alle altre.
- Redigi contenuto concreto suggerito dal titolo, senza preamboli inutili.
- Non scrivere e non ripetere mai '{sezione}' come intestazione. Inizia direttamente con il contenuto.
- Usa formattazione editoriale pulita: non usare Markdown, simboli ###, ##, **, __, ``` o intestazioni tecniche. Se servono elenchi, usa semplici punti o numeri senza caratteri decorativi.
- Non inserire URL, link, citazioni, note bibliografiche o sezioni fonti.
- Se sono disponibili fonti esterne, usale solo per ragionare e integrare concetti pertinenti, senza citarle nel testo finale.
- Prima di consegnare, verifica internamente che il contenuto sia completo per la sezione assegnata, che non sia una bozza o un frammento e che non contenga residui di altre sezioni.

=== PROFONDITÀ ADATTIVA E SPIEGAZIONE PASSO PASSO ===
- Scrivi un testo professionale, completo e proporzionato alla complessità della sezione: amplia dove sono necessari metodo, procedura, decisioni tecniche o esempi; evita invece di allungare con frasi motivazionali, ripetizioni o riassunti inutili.
- Per una procedura o una funzione software, segui questa sequenza: scopo pratico; prerequisiti e strumenti; passaggi numerati nell'ordine esatto; cosa osservare dopo ogni passaggio; controllo del risultato; errori frequenti e correzioni; esempio realistico di applicazione.
- Per un concetto teorico, segui questa sequenza: definizione chiara; perché è importante; come si applica; esempio concreto; limiti, eccezioni o errori da evitare.
- Non saltare passaggi impliciti. Spiega ogni azione in modo che un lettore del livello dichiarato possa ripeterla autonomamente.
- Quando è utile, includi checklist, tabella, caso pratico o criterio verificabile all'interno della spiegazione, senza trasformarli in capitoli separati.
- Mantieni paragrafi leggibili e sottotitoli brevi solo quando migliorano la consultazione; non usare formule generiche come "semplice", "intuitivo" o "fondamentale" senza spiegare concretamente il perché.
"""

def valuta_indice_editoriale(indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti=""):
    """Valuta l'indice rispetto al brief compilato nella sidebar."""
    prompt = f"""Valuta professionalmente il seguente indice editoriale in lingua {lingua}.

DATI DEL BRIEF
Titolo: {titolo}
Trama/argomento: {trama}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}
Obiettivo del libro: {obiettivo}
Approfondimenti prioritari: {approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

INDICE DA VALUTARE
{indice}

Esamina esclusivamente: attinenza al brief, copertura degli approfondimenti prioritari, ordine logico, completezza, progressione del lettore, assenza di ripetizioni, distinzione tra capitoli e sottocapitoli, concretezza dei titoli e capacità di sostenere un libro completo.
Non valutare il libro non ancora scritto e non inventare informazioni mancanti.

Restituisci testo semplice, senza Markdown e senza URL, in questo formato:
VOTO COMPLESSIVO: X/10
VERDETTO: una frase chiara.
PUNTI DI FORZA: 3-5 osservazioni concrete.
MIGLIORAMENTI CONSIGLIATI: soltanto modifiche necessarie e direttamente applicabili; se l'indice è già valido, scrivi "Nessuna modifica strutturale necessaria".
COERENZA CON IL BRIEF: breve verifica di titolo, pubblico, obiettivo, genere e stile.
"""
    return pulisci_testo_editoriale(chiedi_gpt(
        prompt,
        "Sei un editor senior specializzato in architettura di libri. Sei rigoroso, concreto e non usi valutazioni vaghe."
    ))

def firma_controllo_coerenza(indice, contenuti, titolo, trama, genere, stile, narrativa, pov, obiettivo, approfondimenti):
    """Identifica con certezza la versione del manoscritto su cui è stato prodotto il report."""
    parti = [indice, titolo, trama, genere, stile, narrativa, pov, obiettivo, approfondimenti]
    parti.extend(f"{sezione}\n{contenuto}" for sezione, contenuto in contenuti.items())
    return hashlib.sha256("\n\u241e\n".join(str(p or "") for p in parti).encode("utf-8")).hexdigest()


def blocchi_per_audit_manoscritto(contenuti, limite_caratteri=18000):
    """Divide il testo intero in blocchi consecutivi, senza omettere la parte centrale delle sezioni."""
    blocchi, corrente = [], ""
    for sezione, contenuto in contenuti.items():
        testo = pulisci_testo_editoriale(contenuto).strip()
        if not testo:
            continue
        unita = f"SEZIONE: {sezione}\nTESTO:\n{testo}\n\n"
        while unita:
            spazio = limite_caratteri - len(corrente)
            if spazio <= 300:
                blocchi.append(corrente)
                corrente, spazio = "", limite_caratteri
            if len(unita) <= spazio:
                corrente += unita
                unita = ""
            else:
                punto_taglio = unita.rfind("\n", 0, spazio)
                if punto_taglio < max(500, spazio // 2):
                    punto_taglio = spazio
                corrente += unita[:punto_taglio]
                blocchi.append(corrente)
                corrente, unita = "", unita[punto_taglio:]
    if corrente.strip():
        blocchi.append(corrente)
    return blocchi


def chiedi_audit_editoriale(prompt):
    """Esegue l'audit editoriale completo con il modello rapido dedicato ai controlli."""
    try:
        risposta = client.responses.create(
            model="gpt-5-mini",
            input=prompt
        )
        testo = getattr(risposta, "output_text", "") or ""
        return pulisci_testo_editoriale(testo).strip()
    except Exception as e:
        return f"ERRORE AUDIT: {str(e)}"


def valuta_manoscritto_completo(indice, contenuti, titolo, trama, genere, stile, narrativa, pov, obiettivo, lingua, approfondimenti=""):
    """Valuta ogni porzione del manoscritto, poi crea una sintesi basata sugli esiti effettivi."""
    sezioni_scritte = [s for s, c in contenuti.items() if pulisci_testo_editoriale(c).strip()]
    sezioni_vuote = [s for s, c in contenuti.items() if not pulisci_testo_editoriale(c).strip()]
    if not sezioni_scritte:
        return "VALUTAZIONE NON ESEGUIBILE: non è ancora presente testo da analizzare. Genera almeno una sezione e riprova."

    brief = f"""Titolo: {titolo}
Trama/argomento: {trama}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}
Obiettivo: {obiettivo}
Approfondimenti prioritari: {approfondimenti.strip() or "Nessuno"}"""
    blocchi = blocchi_per_audit_manoscritto(contenuti)
    esiti_blocchi = []
    for numero, blocco in enumerate(blocchi, 1):
        esito = chiedi_audit_editoriale(f"""Sei un editor rigoroso. Analizza il BLOCCO {numero} di {len(blocchi)} di un manoscritto in lingua {lingua}.

BRIEF
{brief}

INDICE
{indice}

BLOCCO DA VALUTARE
{blocco}

Valuta soltanto le prove contenute nel blocco: aderenza al titolo e al brief, pertinenza, profondità, chiarezza, progressione locale, eventuali ripetizioni e istruzioni mancanti. Non inventare difetti, non fare verifiche online e non citare fonti. Restituisci testo semplice, senza Markdown, con queste etichette: SEZIONI ESAMINATE, PUNTI FORTI, PROBLEMI SPECIFICI, INTERVENTI PROPOSTI.""")
        esiti_blocchi.append(f"AUDIT BLOCCO {numero}\n{esito}")

    audit_compilati = "\n\n".join(esiti_blocchi)
    sintesi = chiedi_audit_editoriale(f"""Sei un direttore editoriale. Prepara la valutazione finale di un manoscritto in lingua {lingua}, basandoti esclusivamente sul brief, sui dati oggettivi e sugli audit qui sotto. Non inventare contenuti non riportati.

BRIEF
{brief}

DATI OGGETTIVI
Sezioni scritte: {len(sezioni_scritte)}
Sezioni non scritte: {len(sezioni_vuote)}
Elenco sezioni non scritte: {', '.join(sezioni_vuote) if sezioni_vuote else 'nessuna'}

AUDIT DI TUTTI I BLOCCHI DEL MANOSCRITTO
{audit_compilati}

Se ci sono sezioni non scritte, dichiara con chiarezza che la valutazione riguarda un manoscritto parziale e non assegnare un voto finale al libro completo. Altrimenti assegna un voto da 1 a 10.
Non inserire URL, fonti, Markdown o ragionamenti interni. Restituisci esattamente queste voci:
STATO DEL MANOSCRITTO:
VOTO COMPLESSIVO DEL LIBRO:
VERDETTO EDITORIALE:
ADERENZA ALLA SIDEBAR:
PROFONDITÀ DELLE TEMATICHE:
COERENZA E PROGRESSIONE:
SEZIONI DA MIGLIORARE:
AZIONI PRIORITARIE:""")
    return sintesi

def costruisci_specifica_editoriale(titolo, genere, stile, narrativa, pov, obiettivo, argomento, approfondimenti=""):
    """Crea una specifica strutturata che indice e capitoli possono applicare in modo coerente."""
    return f"""=== SPECIFICA EDITORIALE STRUTTURATA ===
Titolo: {titolo}
Genere: {genere}
Tipologia di scrittura: {stile}
Stile di racconto: {narrativa}
Punto di vista: {pov}

OBIETTIVO OPERATIVO:
{obiettivo}

ARGOMENTO E CONFINI:
{argomento}

APPROFONDIMENTI PRIORITARI (FACOLTATIVI):
{approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

Per ogni sezione ricava un risultato concreto, il livello del lettore, i concetti necessari,
gli esempi o le procedure da produrre e ciò che deve restare fuori per evitare ripetizioni.
"""


def audit_simulazioni_test_prep(indice, contenuti, obiettivo, argomento):
    """Controlla che le simulazioni di un test prep contengano davvero quesiti e soluzioni completi."""
    risultati = []
    attese = numero_domande_simulazione(indice, argomento, obiettivo)
    sezioni = [
        sezione for sezione in contenuti
        if sezione_simulazione_test_prep(sezione, indice, "Test Prep (Preparazione Esami)")
    ]
    if not sezioni:
        return ["ERRORE TEST PREP: nell'indice non è stata individuata una sezione che contenga i quesiti di una simulazione."]
    if not attese:
        risultati.append("AVVISO TEST PREP: il brief non dichiara quante domande deve contenere ogni simulazione; il controllo non può validarne la completezza numerica.")

    domande_fuori_simulazione = []
    for sezione, testo in contenuti.items():
        if sezione not in sezioni:
            domande_fuori_simulazione.extend(domande_normalizzate_test_prep(testo))

    for sezione in sezioni:
        testo = contenuti.get(sezione, "")
        domande = domande_normalizzate_test_prep(testo)
        soluzioni = len(re.findall(r"(?im)^\s*soluzione\s+\d{1,3}\s*[:.-]", testo or ""))
        risultati.append(f"Simulazione '{sezione}': domande rilevate {len(domande)}; soluzioni rilevate {soluzioni}.")
        if attese and len(domande) != attese:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' richiede {attese} domande ma ne contiene {len(domande)}.")
        if attese and soluzioni != attese:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' richiede {attese} soluzioni commentate ma ne contiene {soluzioni}.")
        ripetute_interne = len(domande) - len(set(domande))
        ripetute_esterne = len(set(domande) & set(domande_fuori_simulazione))
        if ripetute_interne:
            risultati.append(f"ERRORE TEST PREP: '{sezione}' contiene {ripetute_interne} domande duplicate al suo interno.")
        if ripetute_esterne:
            risultati.append(f"AVVISO TEST PREP: '{sezione}' riutilizza {ripetute_esterne} domande già presenti in quiz o altre sezioni.")
        if (not attese or (len(domande) == attese and soluzioni == attese)) and not ripetute_interne:
            risultati.append(f"OK TEST PREP: struttura numerica della simulazione '{sezione}' valida.")
    return risultati

def analizza_coerenza_libro(indice, contenuti, obiettivo, argomento, genere=""):
    """Controllo deterministico preliminare su struttura, copertura e ripetizioni."""
    risultati = ["REPORT CONTROLLO COERENZA DEL LIBRO"]
    testo = "\n".join(contenuti.values()) if contenuti else ""
    capitoli = re.findall(r"(?im)^(?:Capitolo|Chapter|CAPITOLO)\\s+\\d+", indice or "")
    sottocapitoli = re.findall(r"(?m)^\\d+\\.\\d+\\s+", indice or "")
    risultati.append(f"Capitoli rilevati: {len(capitoli)}")
    risultati.append(f"Sottocapitoli rilevati: {len(sottocapitoli)}")
    if not indice.strip(): risultati.append("ERRORE: indice assente")
    if not obiettivo.strip(): risultati.append("AVVISO: obiettivo assente")
    if not argomento.strip(): risultati.append("AVVISO: argomento assente")
    sezioni_indice = st.session_state.get("lista_capitoli", [])
    if sezioni_indice:
        mancanti = sezioni_mancanti_per_esportazione(sezioni_indice, genere)
        risultati.append(f"Sezioni dell'indice non ancora complete: {len(mancanti)}")
        if mancanti:
            anteprima = "; ".join(mancanti[:8])
            risultati.append(f"ERRORE: non esportare il libro. Sezioni incomplete: {anteprima}" + (" ..." if len(mancanti) > 8 else ""))
        else:
            risultati.append("OK: tutte le sezioni previste dall'indice risultano compilate.")
    if genere == "Ricettario":
        numero_richiesto = estrai_numero_ricette("", argomento, obiettivo)
        capitoli_ricetta = [s for s in sezioni_indice if tipo_sezione_editoriale(s) == "capitolo"]
        if numero_richiesto:
            risultati.append(f"Ricette richieste dal brief: {numero_richiesto}; capitoli-ricetta nell'indice: {len(capitoli_ricetta)}")
            if len(capitoli_ricetta) != numero_richiesto:
                risultati.append("ERRORE RICETTARIO: il numero delle ricette nell'indice non coincide con la richiesta.")
        termini_non_vegani = r"\b(miele|latte vaccino|burro|uova|formaggio|panna|yogurt greco)\b"
        non_vegani = sorted(set(re.findall(termini_non_vegani, testo.lower())))
        if non_vegani:
            risultati.append("AVVISO RICETTARIO: verificare ingredienti non vegani rilevati: " + ", ".join(non_vegani))
    if genere == "Test Prep (Preparazione Esami)":
        risultati.extend(audit_simulazioni_test_prep(indice, contenuti, obiettivo, argomento))
    if len(testo.strip()) < 1000: risultati.append("AVVISO: contenuto ancora troppo breve per una verifica completa")
    frasi = [f.strip().lower() for f in re.split(r"[.!?]+", testo) if len(f.strip()) > 40]
    duplicati = len(frasi) - len(set(frasi))
    risultati.append(f"Frasi duplicate identiche rilevate: {max(0, duplicati)}")
    if duplicati == 0: risultati.append("OK: nessuna duplicazione identica rilevata nel testo disponibile")
    return "\n".join(risultati)

# ======================================================================================================================
# 8. UI PRINCIPALE & GENERAZIONE PROMPT DINAMICO
# ======================================================================================================================
st.markdown(f'<div class="custom-title">AI di Antonino: {val_titolo if val_titolo else "Ebook Creator PRO"}</div>', unsafe_allow_html=True)

sync_capitoli()
lista_cap_base = st.session_state.get("lista_capitoli", [])
opzioni_editor = [L["preface"]] + lista_cap_base + [L["ack"]]

if val_titolo and val_trama:
    
    # VALUTAZIONE DINAMICA: L'IA decide se usare o meno la manipolazione cerebrale
    usa_tre_cervelli = valuta_approccio_neurologico(val_genere, val_stile, val_narrativa)
    
    if usa_tre_cervelli:
        modulo_stilistico = """
=== METODOLOGIA DEI 3 CERVELLI (NEUROMARKETING) ===
Devi strutturare il testo per comunicare simultaneamente con i 3 livelli cerebrali del lettore, iniettando la giusta chimica:
1. CERVELLO RETTILE (Sopravvivenza & Istinto): Usa un linguaggio netto, tangibile e basato sui contrasti (prima/dopo, problema/soluzione). Attira l'attenzione istantaneamente. Elimina parole deboli o passive.
2. CERVELLO LIMBICO (Emozione & Chimica): Usa "Storytelling" ed empatia. Scegli vocaboli sensoriali che stimolino il rilascio di dopamina (curiosità/ricompensa) e ossitocina (fiducia/connessione). Fai percepire al lettore che comprendi esattamente il suo stato d'animo.
3. NEOCORTECCIA (Logica & Dati): Fornisci struttura, dati precisi, ragionamenti logici e prove che giustifichino razionalmente le emozioni suscitate dal sistema limbico.
"""
    else:
        modulo_stilistico = """
=== APPROCCIO ANALITICO E OGGETTIVO ===
Il genere e lo stile scelti richiedono un approccio neutrale e rigoroso. 
NON utilizzare manipolazioni emotive o neuromarketing. Mantieni un tono accademico, logico e fattuale. 
Fornisci dati, structures deduttive e un linguaggio pulito, tipico delle pubblicazioni di alto rigore tecnico-scientifico.
"""

    modulo_fonti = ""
    if st.session_state.get("conoscenza_extra"):
        modulo_fonti = """
=== INTEGRAZIONE FONTI ESTERNE (RAGIONAMENTO AI) ===
I documenti forniti serve per arricchire il tuo ragionamento, estrarre dati e terminologia tecnica.
È TASSATIVAMENTE VIETATO FARE COPIA E INCOLLA dei testi originali. Usa queste fonti esclusivamente come "cervello esterno" per scrivere le tue sezioni originali basandoti su quei concetti, con lo stile e il POV richiesto per il libro.

=== STUDIO ATTIVO E ASSIMILAZIONE CONCETTI ===
Devi agire come uno studioso che ha appena letto le fonti caricate dall'utente. 
1. ESTRAZIONE E ANALISI: Individua i princìpi cardine, i framework, i concetti chiave e i dati presenti nei documenti.
2. RAGIONAMENTO PROFONDO: Non limitarti a citare i concetti a pappagallo. Sviscerali, spiegane il "perché", il contesto e come si applicano operativamente all'argomento del libro.
3. ELABORAZIONE ORIGINALE: Fai tuoi questi concetti. Intrecciali fluidamente con la tua base di conoscenza per creare un testo ricco e autorevole, dimostrando assoluta padronanza della materia, mantenendo il divieto di copia-incolla.
"""

    # --- INIZIO NUOVE RIGHE PER ADATTAMENTO PROMPT IN BASE AL GENERE ---
    modulo_approfondimento_genere = ""
    # Aggiunto "Storico" a questo blocco logico
    if "Manuale" in val_genere or "Saggio" in val_genere or "Test Prep" in val_genere or "Economia" in val_genere or "Storico" in val_genere:
        modulo_approfondimento_genere = """
=== DIRETTIVA DI APPROFONDIMENTO ESTREMO (MANUALISTICA E SAGGISTICA) ===
Trattandosi di un testo tecnico, didattico o manualistico, il tuo compito primario è ISTRUIRE. 
Ogni singolo capitolo e sottocapitolo DEVE essere sviscerato in profondità assoluta. 
- NON dare MAI nulla per scontato o rimanere in superficie.
- Spiega dettagliatamente "COSA è", "PERCHÉ funziona così" e "COME si applica" nella pratica.
- Inserisci esempi concreti, casi d'uso pratici, schemi logici o spiegazioni passo-passo.
- Il lettore deve acquisire una competenza reale, dettagliata e spendibile alla fine di questa sezione. Evita categoricamente la superficialità.
"""
    elif "Romanzo" in val_genere or "Narrativo" in val_genere or "Thriller" in val_genere or "Fantasy" in val_genere or "Fantascienza" in val_genere:
        modulo_approfondimento_genere = """
=== DIRETTIVA DI IMMERSIONE NARRATIVA (NARRATIVA E ROMANZO) ===
Trattandosi di un'opera narrativa, il tuo focus esclusivo è lo STORYTELLING e l'immersione.
- Mostra, non raccontare (Show, Don't Tell). Descrivi minuziosamente gli ambienti, le espressioni e le atmosfere usando i 5 sensi.
- Dai spessore profondo ai personaggi, ai dialoghi e gestisci il ritmo dell'azione o dell'introspezione.
- Evita totalmente lo stile accademico, saggistico o manualistico: il lettore deve "vivere" la scena in tempo reale, non studiarla.
"""
    # --- FINE NUOVE RIGHE ---

    # --- INIZIO NUOVE RIGHE PER APPROFONDIMENTO ED ESEMPI SPECIFICI PER GENERE ---
    modulo_esempi_specifici = """
=== DIRETTIVA DI PROFONDITÀ ED ESEMPI PRATICI (CALIBRATA SUL GENERE) ===
Il tuo compito non è solo descrivere, ma DIMOSTRARE e APPROFONDIRE. In base al genere selezionato, devi generare contenuti operativi o immersivi reali:
"""
    if "Test Prep" in val_genere or "Quiz" in val_genere:
        modulo_esempi_specifici += "- Crea VERE SIMULAZIONI D'ESAME o test complessi inerenti al capitolo. Inserisci domande a risposta multipla, scenari pratici e fornisci le soluzioni dettagliate con spiegazione logica per ogni opzione (perché è giusta o sbagliata).\n"
    elif "Manuale" in val_genere or "Business" in val_genere or "Economia" in val_genere or "Self-Help" in val_genere:
        modulo_esempi_specifici += "- Inserisci veri e propri CASI STUDIO (reali o verosimili), framework applicativi, checklist e scenari di 'Roleplay' o 'What-if' per mostrare come applicare la teoria nella realtà.\n"
    # Aggiunto blocco logico per "Storico"
    elif "Storico" in val_genere:
        modulo_esempi_specifici += "- Inserisci riferimenti storici accurati, date precise, eventi chiave, analisi del contesto socio-politico e cita documenti o figure di rilievo dell'epoca.\n"
    elif "Saggio" in val_genere or "Tecnico" in val_genere:
        modulo_esempi_specifici += "- Fornisci spiegazioni tecniche microscopiche, formule, dati statistici ed esempi concreti di applicazione nel mondo reale per supportare la tesi.\n"
    elif "Romanzo" in val_genere or "Narrativo" in val_genere or "Fantasy" in val_genere or "Thriller" in val_genere:
        modulo_esempi_specifici += "- Crea scene vissute. Mostra interazioni specifiche, dialoghi autentici tra personaggi e reazioni ambientali. Non riassumere gli eventi, falli accadere 'in camera'.\n"
    elif "Ricettario" in val_genere:
        modulo_esempi_specifici += "- Aggiungi varianti degli ingredienti, trucchi dello chef per rimediare agli errori comuni e cenni scientifici sul perché avvengono certe reazioni in cottura.\n"
    else:
        modulo_esempi_specifici += "- Non rimanere mai in superficie: ogni volta che introduci un concetto, fai subito un ESEMPIO PRATICO e dettagliato che lo esplichi al 100%.\n"
    # --- FINE NUOVE RIGHE ---

    # --- INIZIO NUOVE RIGHE PER ADERENZA TITOLO-GENERE (DETTAGLIO ESTREMO) ---
    modulo_aderenza_titolo_genere = f"""
=== DIRETTIVA DI ESECUZIONE REALE E SPECIFICA (INCROCIO GENERE E TITOLO SEZIONE) ===
Analizza attentamente il genere letterario ('{val_genere}') e il titolo esatto della sezione che stai per scrivere.
Non limitarti a fare un discorso generico: devi FORNIRE MATERIALMENTE ciò che il titolo della sezione suggerisce, declinato per quel genere.
- Esempio 1: Se il genere è "Test Prep" o simile e il titolo della sezione che stai scrivendo parla di "Test", "Quiz", "Simulazione" o "Autovalutazione", DEVI generare un VERO test (es. domande a risposta multipla, scenari, soluzioni e spiegazioni dettagliate del perché un'opzione è corretta o sbagliata). Non descrivere come si fa un test, FALLO E REDIGILO REALMENTE.
- Esempio 2: Se il genere è "Manuale", "Business" o "Economia" e il titolo suggerisce un "Piano d'Azione", un "Caso Studio" o un "Esercizio", scrivi i passaggi operativi completi o il caso studio di esempio con nomi, dati, calcoli finanziari e soluzioni.
- Esempio 3: Se il genere è "Ricettario" e la sezione è un piatto, scrivi la ricetta vera e propria.
- Regola Universale: Cogli l'intento pratico implicito nel titolo della sezione corrente in base al genere. Se la sezione richiede un contenuto specifico (una tabella, un test di autovalutazione, un esercizio, un dialogo d'esempio), REDIGI QUEL CONTENUTO FISICAMENTE. Il lettore deve trovarsi di fronte allo strumento o alla scena reale, non a un riassunto teorico.
"""
    # --- FINE NUOVE RIGHE ---

    modulo_operativita_universale = """
=== DIRETTIVA OBBLIGATORIA DI OPERATIVITÀ E DETTAGLIO ===
Ogni sottocapitolo deve insegnare concretamente l'argomento, non limitarsi a descriverlo.
Quando è pertinente, includi sempre: obiettivo, prerequisiti, strumenti, impostazioni,
procedura numerata passo passo, valori o parametri, risultato atteso, errori frequenti,
correzioni, esercizio pratico e criterio verificabile di completamento.
Se la sezione tratta prompt, scrivi prompt completi e copiabili. Se tratta software,
indica menu, comandi, pulsanti e sequenza esatta. Se tratta una procedura, eseguila
materialmente nel testo con dati o esempi realistici. Distingui fatti, esempi ipotetici
e indicazioni da verificare. Non promettere risultati garantiti e non usare formule vaghe.
Nel capitolo principale mantieni la visione d'insieme; nel sottocapitolo sviluppa il
dettaglio assegnato senza anticipare o ripetere gli altri sottocapitoli.
"""

    profilo_genere_corrente = profilo_genere_stesura(val_genere)
    profilo_tipologia_corrente = profilo_tipologia_stesura(val_stile)
    profilo_indice_corrente = profilo_struttura_indice(val_genere, val_titolo, val_trama, val_goal)

    # PROMPT POTENZIATO CON COERENZA POV, PULIZIA SINTATTICA E CONFORMITA' DI GENERE
    S_PROMPT = f"""
Sei un esperto Madrelingua in {lingua_sel}, Editor e Luminare mondiale nel campo '{val_genere}'. 
Stai redigendo l'ebook '{val_titolo}'. 

{modulo_fonti}

{modulo_approfondimento_genere}
{modulo_esempi_specifici}
{modulo_aderenza_titolo_genere}
{modulo_operativita_universale}

=== PROFILO EDITORIALE SPECIFICO ===
- Regole del genere '{val_genere}': {profilo_genere_corrente}
- Regole della tipologia '{val_stile}': {profilo_tipologia_corrente}
- Architettura richiesta: {profilo_indice_corrente}
Queste regole prevalgono sulle istruzioni generiche incompatibili con il genere. Non applicare
procedure, checklist, test, sottocapitoli o scene narrative quando non sono pertinenti.

PARAMETRI DI BASE (DA APPLICARE TASSATIVAMENTE IN OGNI SEZIONE):
- Stile di Racconto: {val_narrativa}
- Obiettivo Emozionale/Pratico: {val_goal}
- Tipologia di Scrittura: {val_stile}
- Punto di Vista (Relazione con il lettore): {val_pov}. Adatta coerentemente questo pronome alla grammatica della lingua {lingua_sel}.
- Conformità di Genere: Il testo DEVE rispecchiare in pieno le regole, la formattazione e la terminologia del genere '{val_genere}' (es. se è un ricettario, usa formati strutturati con ingredienti e step; se è un romanzo usa narrazione fluida; se è 'Test Prep', usa schemi, riassunti puntati, concetti chiave da memorizzare e simulazioni d'esame).
- Lingua di Output Categorica: {lingua_sel}

{specifica_editoriale}

{modulo_stilistico}

=== REGOLA DI FORMATTAZIONE E SINTASSI PULITA (CRITICO) ===
- Usa ESCLUSIVAMENTE una punteggiatura standard, tipografica e impeccabile. 
- SONO SEVERAMENTE VIETATE punteggiature anomale, artefatti markdown inutili, asterischi eccessivi, o emoji nel corpo del testo.
- Il testo deve scorrere con l'eleganza formale e la pulizia di un vero libro stampato (sintassi corretta, paragrafi chiari).

=== REGOLA AUREA: GERARCHIA E NON-RIPETIZIONE (CAPITOLO VS SOTTOCAPITOLO) ===
Dovrai analizzare l'indice fornito per capire la tua esatta posizione:
- SE STAI SCRIVENDO UN CAPITOLO PRINCIPALE (es. 1, 2, 3): Focalizzati sulla visione d'insieme, introduci l'argomento in modo macroscopico. NON rubare i dettagli tecnici, gli esempi specifici o i casi studio che appartengono ai tuoi sottocapitoli.
- SE STAI SCRIVENDO UN SOTTOCAPITOLO (es. 1.1, 1.2, 3.4): Entra inmediatamente nel dettaglio estremo, nell'azione pratica o nell'analisi profonda. NON ripetere mai le premesse o le introduzioni generali già spiegate nel capitolo padre. 
- MEMORIA GLOBALE: Leggi il contesto fornito. Non ripetere mai concetti, parole chiave o aneddoti già utilizzati in altre sezioni.

=== DIRETTIVA ANTI-RIPETIZIONE E BLACKLIST DEGLI ARGOMENTI ===
Il sistema anti-ripetizione è il parametro più critico di questa operazione:
1. DISTINZIONE PADRE/FIGLIO: Se stai scrivendo un Capitolo Principale (es. "Capitolo 1"), devi limitarti a una visione "dall'alto", introducendo i temi SENZA svelarne le meccaniche o gli esempi. Se stai scrivendo un Sottocapitolo (es. "1.1" o "1.2"), devi entrare nel micro-dettaglio e ti è SEVERAMENTE VIETATO riassumere o ripetere l'introduzione già fatta nel Capitolo Padre.
2. BLACKLIST DEI CONTENUTI PRECEDENTI: I contenuti presenti nella "MEMORIA CONTENUTI PRECEDENTI" sono da considerarsi in una BLACKLIST. Non usare MAI le stesse introduzioni, non riciclare esempi e non riproporre gli stessi concetti o checklist. Ogni sezione deve essere 100% inedita rispetto alle precedenti.

=== SILENZIO STAMPA ASSOLUTO SUI SOTTOCAPITOLI (MUTUAMENTE ESCLUSIVI) ===
Questa è la regola d'oro per evitare sovrapposizioni e non farti trattare lo stesso argomento due volte:
1. IL CAPITOLO PARLA DEL "PERCHÉ": Se l'indice ti posiziona nella stesura di un Capitolo Padre (es. "Capitolo 2"), il tuo UNICO compito è creare la cornice concettuale. Ti è IMPOSTO IL SILENZIO STAMPA su qualsiasi argomento, tecnica o dettaglio che abbia un Sottocapitolo dedicato (es. 2.1, 2.2). NON SPIEGARE NIENTE DI SPECIFICO NEL CAPITOLO PADRE.
2. IL SOTTOCAPITOLO PARLA DEL "COME" e del "COSA": Se la sezione è un Sottocapitolo (es. "2.1 L'argomento X"), l'intera spiegazione dell'Argomento X DEVE avvenire ESCLUSIVAMENTE lì. Nel Capitolo Padre, X non doveva essere spiegato, ma al massimo accennato come un titolo nel futuro.
3. CONTROLLO FINALE PRIMA DI GENERARE: Guarda la lista completa dei tuoi sottocapitoli e chiediti: "Sto spiegando in questo testo qualcosa che l'indice dice di spiegare nel prossimo paragrafo numerato?". Se la risposta è SÌ, CANCELLA e astieniti. Lascia vuoto informativo per permettere al Sottocapitolo di esistere senza ripetizioni.

=== DIVIETO DI ANTICIPAZIONE (SPOILER SUI SOTTOCAPITOLI) ===
ASCOLTA ATTENTAMENTE: Se l'indice prevede che un argomento specifico venga trattato in un Sottocapitolo (es. 1.1, 1.2, 1.3), è ASSOLUTAMENTE VIETATO parlarne, menzionarlo o spiegarlo nel Capitolo Padre (es. Capitolo 1).
Il Capitolo Padre deve fungere SOLO da cornice introduttiva generale. Non deve MAI svuotare di significato i sottocapitoli anticipandone i contenuti. Mantieni il vuoto informativo sulle questioni specifiche finché non arrivi a scrivere il sottocapitolo dedicato.

=== APPLICAZIONE DIRETTIVE (STESURA PULITA) ===
Devi interiorizzare e applicare alla lettera le seguenti istruzioni prima di generare il testo:
1. Il genere '{val_genere}'
2. La tipologia di scrittura '{val_stile}' e lo stile di racconto '{val_narrativa}'
3. Il POV '{val_pov}'
4. L'obiettivo '{val_goal}'
CRITICO: NON inserire alcun "ragionamento editoriale", commento, introduzione o meta-testo. L'output DEVE contenere ESCLUSIVAMENTE il contenuto finale del capitolo/sottocapitolo, pronto per la pubblicazione.

=== DIRETTIVA DI CONFORMITÀ ASSOLUTA (PUNTO DI VISTA E STILE) ===
È TASSATIVO e NON NEGOZIABILE che l'intero testo sia redatto utilizzando ESATTAMENTE il Punto di Vista (POV) impostato nella sidebar: "{val_pov}". 
- Se è impostato su "Tu", rivolgiti direttamente e informalmente al singolo lettore (es. "scoprirai che...").
- Se è impostato su "Voi", rivolgiti in modo plurale e autorevole (es. "scoprirete che...").
- Se è impostato su "Noi", usa un approccio inclusivo (es. "scopriremo che...").
- Se è "Impersonale", usa forme impersonali o passive, distaccate e oggettive (es. "si scoprirà che...").
L'intelligenza artificiale DEVE effettuare un controllo lessicale e grammaticale ad ogni fine paragrafo per assicurarsi che non ci siano "scivoloni" o cambi di pronome accidentali. Lo stile di scrittura "{val_stile}" deve permeare ogni singola scelta di vocabolario.

=== REGOLA DELLA DENSITÀ E APPROCCIO DIRETTO (NO FLUFF) ===
- VAI AL SODO: Elimina qualsiasi preambolo inutile, frasi fatte o giri di parole. Inizia immediatamente a trattare il cuore dell'argomento della sezione.
- ZERO VAGHEZZA: Sii estremamente descrittivo, specifico e dettagliato. Non limitarti a enunciare i concetti, ma sviscerali e dimostrali.
- PROFONDITÀ ARGOMENTATIVA: Tratta gli argomenti in maniera fortemente argomentativa. Se stai spiegando una teoria, una tecnica o un concetto pratico, fornisci il "come" e il "perché" con autorevolezza, supportando le tue affermazioni con logica ferrea, dati e dettagli concreti, mantenendo un focus laser sull'argomento.

=== OUTPUT OBBLIGATORI PER EVITARE GENERICITÀ ===
- Ogni sezione deve contenere almeno un elemento applicabile: procedura numerata, esempio concreto,
  checklist, tabella, esercizio, caso studio o criterio di verifica, in base al titolo della sezione.
- Non dichiarare soltanto che una tecnica è utile: mostra quando usarla, come applicarla e come controllare il risultato.
- Non introdurre strumenti, dati o risultati non supportati dalle fonti o dal contesto; segnala ciò che deve essere verificato.
- Rispetta i confini della sezione e non riempire spazio con ripetizioni o frasi motivazionali generiche.

=== APPROCCIO IPER-PRATICO E MICRO-DETTAGLIO ===
- OPERATIVITÀ IMMEDIATA: Spiega esattamente "COME" fare le cose. Inserisci step operativi, checklist, esempi concreti, casi studio reali o template applicativi.
- IPER-DETTAGLIO: Scendi in profondità nel micro-dettaglio. Se menzioni una tecnica, smontala nei suoi componenti base. Il lettore non deve mai chiedersi "Ok, ma in pratica come si fa?". La risposta deve essere già lì, sviscerata in ogni suo singolo passaggio logico e pratico.

=== DIVIETO ASSOLUTO DI RITRASCRIZIONE TITOLI (CRITICO) ===
- NON RITRASCRIVERE o ripetere MAI il nome del capitolo, del sottocapitolo o della sezione all'interno del testo generato o come intestazione (es. non scrivere mai "Capitolo 1" o "1.1 Introduzione" all'inizio).
- Inizia a scrivere DIRETTAMENTE il corpo del testo. L'applicazione impagina i titoli automaticamente; se tu li scrivi, verrà creato un brutto e fastidioso doppione visivo. Non usare `#` o `##` all'inizio per ripetere il titolo che ti è stato assegnato.

=== MAESTRIA LINGUISTICA E PROFONDITÀ DA LUMINARE (CRITICO) ===
- LIVELLO MADRELINGUA ASSOLUTO: Scrivi in {lingua_sel} con la naturalezza, il ritmo e la ricchezza di vocabolario di un autore locale di altissimo livello. Evita categoricamente frasi robotiche, traduzioni letterali o costrutti tipici dell'IA. Usa le sfumature linguistiche, le metafore e le espressioni idiomatiche proprie della lingua {lingua_sel}.
- COMPETENZA VERTICALE (ESPERTO DEL SETTORE): Comportati come un professionista con 30 anni di esperienza reale in questo esatto argomento. Sii chirurgico nei termini tecnici e fornisci dettagli, aneddoti o concetti avanzati che solo un vero "addetto ai lavori" conoscerebbe.
- NO SUPERFICIALITÀ: Non dare risposte generiche o banali. Ogni paragrafo deve trasudare competenza profonda, spiegando i meccanismi interni, le ragioni nascoste e i dettagli tecnici dell'argomento.
"""

    tabs = st.tabs(L["tabs"] + ["🛠️ 5. Formattazione"])

    # TAB 1: INDICE (CHIRURGIA: FIX SENSO LOGICO E PULIZIA ASSOLUTA DELL'INDICE E CONNESSIONE SARTORIALE)
    with tabs[0]:
        if st.button(L["btn_idx"]):
            with st.spinner("Creazione indice (Neuro-Analisi, Connessione Parametri e Strutturazione Logica in corso)..."):
                
                # --- INIZIO NUOVE RIGHE PER TRADUZIONE TERMINI INDICE ---
                trad_termini = {
                    "Italiano": {"parte": "Parte", "cap": "Capitolo"},
                    "English": {"parte": "Part", "cap": "Chapter"},
                    "Español": {"parte": "Parte", "cap": "Capítulo"},
                    "Français": {"parte": "Partie", "cap": "Chapitre"},
                    "Deutsch": {"parte": "Teil", "cap": "Kapitel"},
                    "Română": {"parte": "Partea", "cap": "Capitolul"},
                    "Русский": {"parte": "Часть", "cap": "Глава"},
                    "العربية": {"parte": "الجزء", "cap": "الفصل"},
                    "中文": {"parte": "部分", "cap": "章节"}
                }
                t_parte = trad_termini.get(lingua_sel, trad_termini["Italiano"])["parte"]
                t_cap = trad_termini.get(lingua_sel, trad_termini["Italiano"])["cap"]
                # --- FINE NUOVE RIGHE ---

                # PROMPT BLINDATO PER L'INDICE: Ora prende in carico TUTTI i parametri della sidebar per coerenza assoluta.
                # E include i termini tradotti (f-string per iniezione variabili)
                prompt_idx = f"""Crea l'indice per il libro '{val_titolo}' rigorosamente in lingua {lingua_sel}. 

PARAMETRI EDITORIALI (L'indice deve essere costruito su misura e strettamente attinente a queste caratteristiche):
- Trama/Argomento Centrale: {val_trama}
- Genere Letterario: {val_genere}
- Tipologia di Scrittura: {val_stile}
- Stile di Racconto: {val_narrativa}
- Punto di Vista: {val_pov}
- Obiettivo Emozionale/Pratico: {val_goal}
- Approfondimenti prioritari: {val_approfondimenti.strip() or "Nessun approfondimento aggiuntivo fornito."}

{specifica_editoriale}

Gli approfondimenti prioritari devono essere considerati prima di distribuire gli altri argomenti nell'indice. Trasformali in capitoli o sottocapitoli soltanto quando sono pertinenti al libro e assegna loro una collocazione logica, senza creare duplicazioni o voci generiche.

=== DIRETTIVA SPECIFICA DELLA TIPOLOGIA DI SCRITTURA ===
Tipologia selezionata: {val_stile}
{direttiva_indice_selezionata}

=== ARCHITETTURA ADATTIVA AL GENERE ===
{profilo_struttura_indice(val_genere, val_titolo, val_trama, val_goal)}

=== REGOLA DI STESURA DEL GENERE ===
{profilo_genere_stesura(val_genere)}

=== SPECIFICA OPERATIVA PER LA PROGETTAZIONE DELL'INDICE ===
Costruisci l'indice come un progetto editoriale eseguibile, non come un elenco generico.
Ricava dal brief il risultato finale promesso, il pubblico e il livello di partenza, i problemi
concreti, il metodo didattico, i deliverable e i limiti del libro.
Definisci una sequenza dal livello iniziale al risultato finale. Ogni Parte deve avere una
funzione distinta; ogni Capitolo deve avere un obiettivo autonomo; ogni sottocapitolo deve
avere un confine preciso, un risultato concreto e almeno un deliverable coerente: procedura,
prompt copiabile, esempio eseguibile, checklist, tabella, esercizio, caso studio o criterio
di verifica. Non creare sottocapitoli ripetitivi.
Distribuisci gli argomenti dell'obiettivo e della trama senza anticipare tutto nell'introduzione.
Per strumenti o software soggetti ad aggiornamento, separa principi stabili, funzioni da verificare
e applicazioni. Mantieni coerenza con genere, tipologia, stile, POV, obiettivo e argomento.
L'indice deve permettere di scrivere sezioni dettagliate senza riempitivi.
"""
                if st.session_state.get("conoscenza_extra"):
                    prompt_idx += f"\n\nFONTI ESTERNE E RAGIONAMENTO:\nUsa queste informazioni fornite dall'utente per strutturare l'indice in modo logico e autorevole. \n{st.session_state['conoscenza_extra'][:4000]}\n"

                prompt_idx += f"""
REGOLE FONDAMENTALI ED ESCLUSIVE:
0. ATTINENZA ASSOLUTA: Inserisci esclusivamente capitoli e sottocapitoli direttamente pertinenti al titolo, alla trama, al pubblico e all'obiettivo del libro. Non aggiungere sezioni generiche o accessorie come glossario dei termini, elenco di risorse, checklist generiche, bibliografia, link, ringraziamenti, conclusioni vaghe o suggerimenti finali. Ogni voce dell'indice deve sviluppare un argomento reale del libro e poter essere trasformata in contenuto sostanziale.
1. SOLO L'INDICE: Non inserire convenevoli, saluti, introduzioni o conclusioni. L'output deve contenere ESCLUSIVAMENTE la lista dell'indice. Nient'altro.
2. COERENZA ASSOLUTA: I titoli dei capitoli e sottocapitoli devono riflettere perfectly lo stile, il genere e la trama richiesta. Se è un ricettario, l'indice deve sembrare un menu; se è un thriller, i capitoli devono creare suspense.
3. ESTENSIONE PROPORZIONATA: Applica l'architettura adattiva indicata sopra. Non imporre 15-18 capitoli, 3-5 sottocapitoli o 100 pagine a generi che richiedono una struttura diversa. Per un ricettario, il numero di Capitoli deve coincidere con il numero di ricette richiesto e ogni Capitolo deve essere una ricetta, senza sottocapitoli artificiosi. Per la narrativa, non frammentare scene in sottocapitoli riempitivi.
4. STRUTTURA GERARCHICA RIGIDA E PULITA: Usa unicamente ed esattamente questo formato di elencazione, SENZA ASTERISCHI O SIMBOLI STRANI:
   {t_parte} I: [Nome Parte]
   {t_cap} 1: [Nome Capitolo]
   1.1 [Sottocapitolo]
   1.2 [Sottocapitolo]
5. SENSO LOGICO SEQUENZIALE: Il flusso narrativo/didattico deve essere ineccepibile. Parti dalle basi/introduzione, sviluppa il cuore del problema, e concludi con soluzioni o risoluzioni finali.
6. PULIZIA VISIVA: Nessuna descrizione sotto i capitoli. Nessuna punteggiatura anomala. Solo l'elenco nudo e crudo.

7. APPLICAZIONE SILENZIOSA DEI PARAMETRI: Applica rigorosamente le istruzioni della sidebar garantendo una perfetta coerenza editoriale. CRITICO: NON inserire alcun "ragionamento strutturale", commento preliminare o spiegazione. Stampa SOLO ed ESCLUSIVAMENTE la lista dell'indice nuda e cruda.

8. PRATICITÀ ESTREMA E IPER-DETTAGLIO: I titoli devono essere estremamente pratici e orientati all'azione. Niente macro-concetti vaghi. Ogni capitolo e sottocapitolo deve puntare a risolvere un problema specifico, mostrando il "come fare" passo dopo passo, con un taglio estremamente operativo e profondo.

9. COMPLETEZZA SENZA RIEMPITIVI: Rispetta il numero e il formato stabiliti dall'architettura adattiva. Ogni Capitolo deve avere una funzione autonoma. Crea sottocapitoli soltanto quando sviluppano aspetti distinti e non quando ripetono ingredienti, procedimenti, esempi o scene già assegnati. Prima di concludere, conta internamente le voci richieste e verifica che nessuna sia vuota o solo un titolo.

10. ADATTAMENTO AL TIPO DI LIBRO E OUTPUT FINALE: Per manuali tecnici separa fondamenti, strumenti, procedure, verifiche e progetto applicativo. Per manuali pratici inserisci esercizi, checklist e risultati misurabili. Per business, marketing, economia e self-help inserisci framework, casi studio, piani d'azione e criteri di valutazione. Per saggi scientifici o storici separa contesto, tesi, prove, fonti e conclusioni. Per ricettari con un numero dichiarato di ricette, ogni Capitolo deve essere una ricetta e non sono ammessi Capitoli introduttivi su tecniche, ingredienti o sicurezza. Per test prep inserisci teoria, esercizi, simulazioni e soluzioni. Per narrativa costruisci sviluppo di trama, personaggi, conflitto e risoluzione, senza imporre procedure tecniche e con titoli di capitolo specifici del brief. In ogni caso prevedi un output finale coerente con il genere: progetto, piano, esercizio completato, ricetta, simulazione, decisione applicativa, sintesi o conclusione narrativa. Gli esempi devono essere concreti e verificabili secondo il tipo di libro.
"""
                
                indice_generato = genera_indice_controllato(
                    prompt_idx, "Senior Book Architect esperto in flow logico-narrativo e design editoriale pulito.",
                    val_genere, val_titolo, val_trama, val_goal, lingua_sel, val_stile, val_narrativa, val_pov
                )
                st.session_state.pop("analisi_voto_indice", None)
                if indice_generato:
                    st.session_state["indice_raw"] = indice_generato
                    sync_capitoli(); st.rerun()
                st.session_state["indice_raw"] = ""
                st.session_state["lista_capitoli"] = []
                st.error(st.session_state.get("ultimo_controllo_indice", "Indice non approvato: riprova con un brief più specifico."))
                
        # FIX ANTI-RESET PER L'INDICE: Salvataggio sicuro per prevenire sovrascritture da parte di Streamlit
        testo_corrente = st.session_state.get("indice_raw", "")
        if st.session_state.get("ultimo_controllo_indice"):
            esito_indice = st.session_state["ultimo_controllo_indice"]
            if "10/10" in esito_indice:
                st.success(esito_indice)
            elif "richiede" in esito_indice or "non ha raggiunto" in esito_indice:
                st.warning(esito_indice)
        testo_input = st.text_area("Indice Gerarchico:", value=testo_corrente, height=400)
        
        if testo_input != testo_corrente:
            # Se la UI ricarica e invia stringa vuota per errore, ignoriamo l'aggiornamento, preservando i dati
            if testo_input.strip() == "" and testo_corrente != "":
                pass
            else:
                st.session_state["indice_raw"] = testo_input
                st.session_state.pop("analisi_voto_indice", None)
                
        if st.button(L["btn_sync"]): sync_capitoli(); st.rerun()

        indice_da_valutare = st.session_state.get("indice_raw", "").strip()
        if indice_da_valutare:
            if st.button("⭐ VOTO INDICE", use_container_width=True, key="voto_indice"):
                with st.spinner("Analisi editoriale dell'indice in corso..."):
                    st.session_state["analisi_voto_indice"] = valuta_indice_editoriale(
                        indice_da_valutare, val_titolo, val_trama, val_genere, val_stile,
                        val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti
                    )
            if st.session_state.get("analisi_voto_indice"):
                st.text_area(
                    "Analisi e voto dell'indice",
                    value=st.session_state["analisi_voto_indice"],
                    height=320,
                    key="output_voto_indice"
                )

    # TAB 2: SCRITTURA E QUIZ (E ORA ANCHE RICETTE)
    with tabs[1]:
        if not lista_cap_base: st.warning(L["msg_err_idx"])
        else:
            # La stesura completa deve seguire esattamente tutte le sezioni disponibili nell'editor:
            # prefazione, parti, capitoli, sottocapitoli e ringraziamenti.
            sezioni_intero_libro = opzioni_editor
            st.caption(f"Stesura completa disponibile: {len(sezioni_intero_libro)} sezioni rilevate. I contenuti già scritti verranno conservati.")
            # Job persistente: genera una sezione alla volta e conserva sempre quanto già scritto.
            # Questo rende la pausa effettiva tra una richiesta AI e la successiva.
            if st.button("📚 SCRIVI TUTTO IL LIBRO", use_container_width=True, key="scrivi_tutto_libro"):
                da_generare = [
                    sezione for sezione in sezioni_intero_libro
                    if not st.session_state.get(chiave_sezione(sezione), "").strip()
                ]
                if not da_generare:
                    st.info("Il libro risulta già scritto: nessun contenuto è stato sovrascritto.")
                else:
                    st.session_state["job_scrittura_coda"] = da_generare
                    st.session_state["job_scrittura_totale"] = len(da_generare)
                    st.session_state["job_scrittura_attivo"] = True
                    st.session_state["job_scrittura_pausa"] = False
                    st.session_state.pop("job_scrittura_errore", None)

            coda_scrittura = st.session_state.get("job_scrittura_coda", [])
            if st.session_state.get("job_scrittura_attivo") and coda_scrittura:
                totale = st.session_state.get("job_scrittura_totale", len(coda_scrittura))
                completati = totale - len(coda_scrittura)
                st.progress(
                    int(completati / totale * 100),
                    text=f"Stesura in corso: completate {completati} di {totale} sezioni."
                )
                sezione_corrente = coda_scrittura[0]
                col_stato, col_pausa = st.columns([3, 1])
                with col_stato:
                    st.info(f"Elaborazione in corso: {sezione_corrente}. Puoi fermare il lavoro prima della sezione successiva.")
                with col_pausa:
                    pausa_richiesta = st.button("⏸ PAUSA", use_container_width=True, key="pausa_scrittura_libro")
                if pausa_richiesta:
                    st.session_state["job_scrittura_attivo"] = False
                    st.session_state["job_scrittura_pausa"] = True
                    st.info(f"Generazione in pausa. Restano {len(coda_scrittura)} sezioni da scrivere; puoi controllare il libro e poi riprendere.")
                else:
                    try:
                        prompt = crea_prompt_stesura_sezione(
                            sezione_corrente, st.session_state['indice_raw'], val_trama, val_genere,
                            val_stile, val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti
                        )
                        st.session_state[chiave_sezione(sezione_corrente)] = genera_contenuto_editoriale(
                            prompt, S_PROMPT, sezione_corrente, st.session_state['indice_raw'], val_trama,
                            val_genere, val_goal, lingua_sel
                        )
                        st.session_state["job_scrittura_coda"] = coda_scrittura[1:]
                        st.rerun()
                    except Exception as exc:
                        st.session_state["job_scrittura_attivo"] = False
                        st.session_state["job_scrittura_pausa"] = True
                        st.session_state["job_scrittura_errore"] = f"{sezione_corrente}: {exc}"
                        st.error("Generazione sospesa per un errore. I contenuti precedenti sono salvi: controllali e poi riprendi.")

            if st.session_state.get("job_scrittura_pausa") and st.session_state.get("job_scrittura_coda"):
                rimanenti = len(st.session_state["job_scrittura_coda"])
                st.warning(f"Generazione in pausa: restano {rimanenti} sezioni. Puoi esaminare l'anteprima e riprendere quando vuoi.")
                if st.session_state.get("job_scrittura_errore"):
                    st.caption(f"Ultimo errore: {st.session_state['job_scrittura_errore']}")
                if st.button("▶ RIPRENDI GENERAZIONE", use_container_width=True, key="riprendi_scrittura_libro"):
                    st.session_state["job_scrittura_attivo"] = True
                    st.session_state["job_scrittura_pausa"] = False
                    st.session_state.pop("job_scrittura_errore", None)
                    st.rerun()
            elif st.session_state.get("job_scrittura_attivo") is False and not st.session_state.get("job_scrittura_coda") and st.session_state.get("job_scrittura_totale"):
                st.success("Libro completato: tutte le sezioni previste sono state generate e salvate.")
            sez_scelta = st.selectbox(L["lbl_sec"], opzioni_editor)
            k_sessione = f"txt_{sez_scelta.replace(' ', '_').replace('.', '')}"
            sottocapitoli_capitolo = individua_sottocapitoli_del_capitolo(sez_scelta, lista_cap_base)
            if sottocapitoli_capitolo:
                chiave_audit_capitolo = f"audit_fatti_{sez_scelta.replace(' ', '_').replace('.', '')}"
                st.caption(f"Capitolo selezionato: verranno elaborati {len(sottocapitoli_capitolo)} sottocapitoli non ancora scritti.")
                if st.button("📝 SCRIVI TUTTI I SOTTOCAPITOLI DEL CAPITOLO", use_container_width=True):
                    da_generare = []
                    gia_presenti = 0
                    for sottocapitolo in sottocapitoli_capitolo:
                        chiave = f"txt_{sottocapitolo.replace(' ', '_').replace('.', '')}"
                        if st.session_state.get(chiave, "").strip():
                            gia_presenti += 1
                        else:
                            da_generare.append((sottocapitolo, chiave))
                    if not da_generare:
                        st.info("Tutti i sottocapitoli di questo capitolo sono già presenti: nessun contenuto è stato sovrascritto.")
                    else:
                        avanzamento = st.progress(0, text="Preparazione della stesura del capitolo...")
                        for posizione, (sottocapitolo, chiave) in enumerate(da_generare, start=1):
                            avanzamento.progress(
                                int((posizione - 1) / len(da_generare) * 100),
                                text=f"Scrittura di {sottocapitolo} ({posizione}/{len(da_generare)})..."
                            )
                            prompt = crea_prompt_stesura_sezione(
                                sottocapitolo, st.session_state['indice_raw'], val_trama, val_genere,
                                val_stile, val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti
                            )
                            st.session_state[chiave] = genera_contenuto_editoriale(
                                prompt, S_PROMPT, sottocapitolo, st.session_state['indice_raw'], val_trama,
                                val_genere, val_goal, lingua_sel
                            )
                        avanzamento.progress(100, text="Sottocapitoli completati.")
                        contenuti_capitolo = [
                            (sottocapitolo, st.session_state.get(f"txt_{sottocapitolo.replace(' ', '_').replace('.', '')}", ""))
                            for sottocapitolo in sottocapitoli_capitolo
                        ]
                        st.session_state[chiave_audit_capitolo] = audit_fatti_capitolo(
                            sez_scelta, contenuti_capitolo, lingua_sel
                        )
                        messaggio = f"Completati {len(da_generare)} sottocapitoli."
                        if gia_presenti:
                            messaggio += f" Conservati senza modifiche: {gia_presenti}."
                        st.success(messaggio)
                        st.rerun()
                if st.button("🔎 CONTROLLA I FATTI DEL CAPITOLO", use_container_width=True):
                    contenuti_capitolo = [
                        (sottocapitolo, st.session_state.get(f"txt_{sottocapitolo.replace(' ', '_').replace('.', '')}", ""))
                        for sottocapitolo in sottocapitoli_capitolo
                    ]
                    with st.spinner("Controllo online mirato dei soli dati aggiornabili del capitolo..."):
                        st.session_state[chiave_audit_capitolo] = audit_fatti_capitolo(
                            sez_scelta, contenuti_capitolo, lingua_sel
                        )
                if st.session_state.get(chiave_audit_capitolo):
                    with st.expander("🔎 Esito del controllo fatti del capitolo", expanded=False):
                        st.write(st.session_state[chiave_audit_capitolo])
            c1, c2, c3 = st.columns([2, 2, 1])
            with c1:
                if st.button(L["btn_write"]):
                    with st.spinner(L["msg_run"]):
                        full_prompt = crea_prompt_stesura_sezione(
                            sez_scelta, st.session_state['indice_raw'], val_trama, val_genere,
                            val_stile, val_narrativa, val_pov, val_goal, lingua_sel, val_approfondimenti
                        )
                        st.session_state[k_sessione] = genera_contenuto_editoriale(
                            full_prompt, S_PROMPT, sez_scelta, st.session_state['indice_raw'], val_trama,
                            val_genere, val_goal, lingua_sel
                        )
            with c2:
                istr = st.text_input(L["btn_edit"], key=f"mod_{k_sessione}", placeholder="Es: Potenzia l'esposizione...")
                if st.button(L["btn_edit"] + " 🪄"):
                    if k_sessione in st.session_state:
                        st.session_state[k_sessione] = pulisci_testo_editoriale(chiedi_gpt(f"Rielabora con focus su: {istr} mantenendo categoricamente la lingua {lingua_sel}, il POV ({val_pov}) e senza usare punteggiatura anomala né riscrivere il titolo all'inizio. Non inserire URL, link, citazioni o sezioni bibliografiche. Testo da modificare:\n{st.session_state[k_sessione]}", S_PROMPT)); st.rerun()
            with c3:
                if st.button("🧠 QUIZ"):
                    if k_sessione in st.session_state:
                        with st.spinner("Generazione Quiz didattico..."):
                            res_q = chiedi_gpt(f"Crea quiz di 10 domande in lingua {lingua_sel} dando del {val_pov} al lettore su:\n{st.session_state[k_sessione]}", "Learning Expert.")
                            st.session_state[k_sessione] += f"\n\nTEST DI VALUTAZIONE\n\n" + pulisci_testo_editoriale(res_q); st.rerun()

                # --- INIZIO NUOVE RIGHE PER TRADUZIONE ESEMPI ---
                trad_esempi = {
                    "Italiano": {"btn": "💡 10 ESEMPI", "titolo": "### 💡 10 ESEMPI PRATICI"},
                    "English": {"btn": "💡 10 EXAMPLES", "titolo": "### 💡 10 PRACTICAL EXAMPLES"},
                    "Español": {"btn": "💡 10 EJEMPLOS", "titolo": "### 💡 10 EJEMPLOS PRÁCTICOS"},
                    "Français": {"btn": "💡 10 EXEMPLES", "titolo": "### 💡 10 EXEMPLES PRATIQUES"},
                    "Deutsch": {"btn": "💡 10 BEISPIELE", "titolo": "### 💡 10 PRAKTISCHE BEISPIELE"},
                    "Română": {"btn": "💡 10 EXEMPLE", "titolo": "### 💡 10 EXEMPLE PRACTICE"},
                    "Русский": {"btn": "💡 10 ПРИМЕРОВ", "titolo": "### 💡 10 ПРАКТИЧЕСКИХ ПРИМЕРОВ"},
                    "العربية": {"btn": "💡 10 أمثلة", "titolo": "### 💡 10 أمثلة عملية"},
                    "中文": {"btn": "💡 10 个例子", "titolo": "### 💡 10 个实际例子"}
                }
                t_btn_ese = trad_esempi.get(lingua_sel, trad_esempi["Italiano"])["btn"]
                t_tit_ese = trad_esempi.get(lingua_sel, trad_esempi["Italiano"])["titolo"]
                # --- FINE NUOVE RIGHE ---

                # --- AGGIUNTA PULSANTE GENERATORE ESEMPI ---
                if st.button(t_btn_ese):
                    if k_sessione in st.session_state:
                        with st.spinner(f"Creazione 10 esempi in {lingua_sel}..."):
                            mem_esempi = st.session_state.get(k_sessione, "")
                            p_esempi = f"""Genera ESATTAMENTE 10 ESEMPI PRATICI, unici e dettagliati rigorosamente in lingua {lingua_sel} per la sezione '{sez_scelta}', perfettamente coerenti con l'argomento: '{val_trama}' e il genere '{val_genere}'.
                            Usa il punto di vista '{val_pov}'.
                            
                            [ATTENZIONE ALLA LINGUA]: È TASSATIVO che l'intero contenuto, inclusi i titoli, sia scritto in {lingua_sel}.
                            
                            [REGOLA ANTI-RIPETIZIONE]: Leggi i contenuti già generati qui sotto e NON RIPETERLI MAI. Crea scenari, casi studio o applicazioni completamente nuovi:
                            
                            {mem_esempi[-4000:]}"""
                            
                            res_e = chiedi_gpt(p_esempi, f"Sei un autorevole esperto in {val_genere} e scrittore in lingua {lingua_sel}.")
                            st.session_state[k_sessione] += f"\n\n{pulisci_testo_editoriale(t_tit_ese)}\n\n" + pulisci_testo_editoriale(res_e)
                            st.rerun()

                # --- INIZIO NUOVE RIGHE PER TRADUZIONE RICETTE ---
                trad_ricette = {
                    "Italiano": {"btn": "🍳 10 RICETTE", "titolo": "### 🍳 10 NUOVE RICETTE", "struttura": "Titolo, Tempi (Preparazione/Cottura), Ingredienti, Procedimento"},
                    "English": {"btn": "🍳 10 RECIPES", "titolo": "### 🍳 10 NEW RECIPES", "struttura": "Title, Prep/Cook Time, Ingredients, Instructions"},
                    "Español": {"btn": "🍳 10 RECETAS", "titolo": "### 🍳 10 NUEVAS RECETAS", "struttura": "Título, Tiempo de preparación/cocción, Ingredientes, Elaboración"},
                    "Français": {"btn": "🍳 10 RECETTES", "titolo": "### 🍳 10 NOUVELLES RECETTES", "struttura": "Titre, Temps de préparation/cuisson, Ingrédients, Préparation"},
                    "Deutsch": {"btn": "🍳 10 REZEPTE", "titolo": "### 🍳 10 NEUE REZEPTE", "struttura": "Titel, Zubereitungs-/Kochzeit, Zutaten, Zubereitung"},
                    "Română": {"btn": "🍳 10 REȚETE", "titolo": "### 🍳 10 REȚETE NOI", "struttura": "Titlu, Timp de preparare/gătire, Ingrediente, Mod de preparare"},
                    "Русский": {"btn": "🍳 10 РЕЦЕПТОВ", "titolo": "### 🍳 10 НОВЫХ РЕЦЕПТОВ", "struttura": "Название, Время подготовки/приготовления, Ингредиенты, Инструкции"},
                    "العربية": {"btn": "🍳 10 وصفات", "titolo": "### 🍳 10 وصفات جديدة", "struttura": "العنوان، وقت التحضير/الطهي، المكونات، طريقة التحضير"},
                    "中文": {"btn": "🍳 10 个食谱", "titolo": "### 🍳 10 个新食谱", "struttura": "标题, 准备/烹饪时间, 配料, 制作步骤"}
                }
                t_btn_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["btn"]
                t_tit_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["titolo"]
                t_strut_ric = trad_ricette.get(lingua_sel, trad_ricette["Italiano"])["struttura"]
                # --- FINE NUOVE RIGHE ---
                
                # --- AGGIUNTA PULSANTE GENERATORE RICETTE ---
                if st.button(t_btn_ric):
                    if k_sessione in st.session_state:
                        with st.spinner(f"Creazione 10 ricette uniche in {lingua_sel}..."):
                            mem_ricette = st.session_state.get(k_sessione, "")
                            p_ricette = f"""Crea ESATTAMENTE 10 RICETTE professionali, uniche e dettagliate rigorosamente in lingua {lingua_sel} per la sezione '{sez_scelta}', perfettamente coerenti con l'argomento: '{val_trama}'.
                            Usa il punto di vista '{val_pov}'.
                            
                            [ATTENZIONE ALLA LINGUA]: È TASSATIVO che l'intera ricetta, inclusi i titoli e le voci strutturali, sia scritta in {lingua_sel}.
                            STRUTTURA DI OGNI RICETTA ({lingua_sel}): {t_strut_ric}. Nessuna emoji.
                            
                            [REGOLA ANTI-RIPETIZIONE ASSOLUTA]: Leggi le ricette o i contenuti già generati qui sotto e NON RIPETERLI MAI. Crea varianti e piatti completamente nuovi:
                            
                            {mem_ricette[-4000:]}"""
                            
                            res_r = chiedi_gpt(p_ricette, f"Sei un autorevole Chef stellato e scrittore di ricettari in lingua {lingua_sel}.")
                            st.session_state[k_sessione] += f"\n\n{pulisci_testo_editoriale(t_tit_ric)}\n\n" + pulisci_testo_editoriale(res_r)
                            st.rerun()

            st.divider()
            st.subheader("🖼️ Inserisci immagine del capitolo")
            st.caption("Crea l'immagine esternamente e caricala qui: verrà inserita nell'anteprima, nel Word e nel PDF della sezione selezionata.")
            file_immagine = st.file_uploader(
                "Carica un'immagine PNG, JPG o WEBP",
                type=["png", "jpg", "jpeg", "webp"],
                key=f"upload_immagine_{k_sessione}"
            )
            if file_immagine:
                img_bytes = normalizza_immagine_caricata(file_immagine)
                if img_bytes:
                    st.session_state.setdefault("immagini_capitoli", {})
                    st.session_state["immagini_capitoli"][sez_scelta] = {
                        "bytes": img_bytes,
                        "caption": f"Immagine: {sez_scelta}",
                        "nome_file": file_immagine.name
                    }
                    st.success(f"Immagine '{file_immagine.name}' associata a: {sez_scelta}.")
            immagine_associata = st.session_state.get("immagini_capitoli", {}).get(sez_scelta)
            if immagine_associata:
                st.image(immagine_associata["bytes"], caption="Immagine associata al capitolo", width=420)

            testo_editor = pulisci_testo_editoriale(st.session_state.get(k_sessione, ""))
            st.session_state[k_sessione] = st.text_area(L["label_editor"], value=testo_editor, height=500)
            
            with st.expander("🔍 Linter Qualità & Analisi Sintattica Avanzata"):
                if st.button("Genera Report Sintattico"): st.write(analizza_qualita_prosa(st.session_state.get(k_sessione, "")))

    # TAB 3: ANTEPRIMA
    with tabs[2]:
        st.subheader(L["preview_tit"])
        contenuti_libro = {
            s: st.session_state.get(f"txt_{s.replace(' ', '_').replace('.', '')}", "")
            for s in opzioni_editor
        }
        firma_attuale_coerenza = firma_controllo_coerenza(
            st.session_state.get("indice_raw", ""), contenuti_libro, val_titolo, val_trama,
            val_genere, val_stile, val_narrativa, val_pov, val_goal, val_approfondimenti
        )
        if st.button("🔍 CONTROLLO COERENZA COMPLETO"):
            with st.spinner("Analisi completa del manoscritto in corso..."):
                controllo_tecnico = analizza_coerenza_libro(
                    st.session_state.get("indice_raw", ""), contenuti_libro, val_goal, val_trama, val_genere
                )
                valutazione_editoriale = valuta_manoscritto_completo(
                    st.session_state.get("indice_raw", ""), contenuti_libro, val_titolo,
                    val_trama, val_genere, val_stile, val_narrativa, val_pov, val_goal,
                    lingua_sel, val_approfondimenti
                )
                st.session_state["report_coerenza_libro"] = (
                    f"CONTROLLO TECNICO\n{controllo_tecnico}\n\n"
                    f"VALUTAZIONE EDITORIALE COMPLETA\n{valutazione_editoriale}"
                )
                st.session_state["report_coerenza_firma"] = firma_attuale_coerenza
        if st.session_state.get("report_coerenza_libro"):
            if st.session_state.get("report_coerenza_firma") != firma_attuale_coerenza:
                st.warning("Analisi non aggiornata: il testo, l'indice o il brief sono cambiati dopo l'ultimo controllo. Premi di nuovo il pulsante per ottenere il report della versione corrente.")
            else:
                st.text_area(
                    "Analisi completa del libro",
                    value=st.session_state["report_coerenza_libro"],
                    height=420,
                    key="output_report_coerenza_libro"
                )
        html_p = f"<div class='preview-box'><h1 style='text-align:center;'>{val_titolo.upper()}</h1>"
        if val_autore: html_p += f"<h3 style='text-align:center;'>di {val_autore}</h3>"
        html_p += "<hr><br>"
        for s in opzioni_editor:
            sk = f"txt_{s.replace(' ', '_').replace('.', '')}"
            if sk in st.session_state and st.session_state[sk].strip():
                html_p += f"<h2>{s.upper()}</h2>"
                img = st.session_state.get("immagini_capitoli", {}).get(s)
                if img:
                    img_b64 = base64.b64encode(img["bytes"]).decode("ascii")
                    caption = img.get("caption", "Immagine didattica")
                    html_p += (
                        f"<div style='text-align:center;margin:18px 0;'>"
                        f"<img src='data:image/png;base64,{img_b64}' "
                        f"style='max-width:58%;height:auto;max-height:360px;object-fit:contain;'>"
                        f"<div style='font-size:13px;color:#555;font-style:italic;'>{caption}</div></div>"
                    )
                testo_preview = pulisci_testo_editoriale(st.session_state[sk])
                html_p += f"<p>{testo_preview.replace(chr(10), '<br>')}</p>"
        st.markdown(html_p + "</div>", unsafe_allow_html=True)

    # TAB 4: ESPORTAZIONE
    with tabs[3]:
        sezioni_incomplete_export = sezioni_mancanti_per_esportazione(lista_cap_base, val_genere)
        if sezioni_incomplete_export:
            st.warning(
                "Esportazione disponibile come BOZZA: alcune sezioni dell'indice sono vuote o troppo brevi. "
                "Il file non è ancora pronto per la pubblicazione."
            )
            st.caption("Sezioni da completare: " + "; ".join(sezioni_incomplete_export[:12]) + (" ..." if len(sezioni_incomplete_export) > 12 else ""))
        else:
            st.success("Controllo completezza superato: tutte le sezioni previste dall'indice sono presenti.")
        cw, cp = st.columns(2)
        with cw:
            if st.button(L["btn_word"]):
                doc = Document(); doc.add_heading(val_titolo, 0)
                if sezioni_incomplete_export:
                    doc.add_paragraph("BOZZA NON COMPLETA - Non pronta per la pubblicazione.")
                for s in opzioni_editor:
                    ke = chiave_sezione(s)
                    if st.session_state.get(ke, "").strip():
                        doc.add_page_break(); doc.add_heading(s.upper(), level=1)
                        img = st.session_state.get("immagini_capitoli", {}).get(s)
                        if img:
                            doc.add_picture(BytesIO(img["bytes"]), width=Inches(4.3))
                            doc.add_paragraph(img.get("caption", ""))
                        doc.add_paragraph(pulisci_testo_editoriale(st.session_state[ke]))
                bw = BytesIO(); doc.save(bw); bw.seek(0)
                suffisso = "_BOZZA" if sezioni_incomplete_export else ""
                st.download_button(L["btn_word"], data=bw, file_name=f"{val_titolo}{suffisso}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with cp:
            if st.button(L["btn_pdf"]):
                pdf = EbookPDF(val_titolo, val_autore); pdf.cover_page()
                if sezioni_incomplete_export:
                    pdf.add_content("BOZZA NON COMPLETA", "Questo file è una bozza di lavoro. Alcune sezioni previste dall'indice non sono ancora state completate; non usarlo per la pubblicazione.")
                for s in opzioni_editor:
                    kd = chiave_sezione(s)
                    if st.session_state.get(kd, "").strip():
                        img = st.session_state.get("immagini_capitoli", {}).get(s)
                        pdf.add_content(
                            s.upper(), pulisci_testo_editoriale(st.session_state[kd]),
                            image_bytes=img.get("bytes") if img else None,
                            image_caption=img.get("caption") if img else None
                        )
                out_p = pdf.output(dest='S').encode('latin-1', 'replace')
                suffisso = "_BOZZA" if sezioni_incomplete_export else ""
                st.download_button(L["btn_pdf"], data=out_p, file_name=f"{val_titolo}{suffisso}.pdf", mime="application/pdf")

    # TAB 5: FORMATTAZIONE E METADATI KDP
    with tabs[4]:
        st.subheader("🛠️ Formattazione")
        st.caption("Carica un manoscritto DOCX o PDF per generare metadati; i file DOCX possono anche essere formattati per il formato KDP 6×9.")
        manoscritto = st.file_uploader(
            "Carica manoscritto",
            type=["docx", "pdf"],
            key="manoscritto_formattazione"
        )
        if manoscritto:
            col_metadati, col_formato = st.columns(2)
            with col_metadati:
                st.markdown("### Metadati KDP")
                lingua_metadati = st.selectbox(
                    "Lingua dei metadati",
                    ["Italiano", "Inglese", "Spagnolo", "Francese", "Tedesco", "Rumeno", "Russo", "Arabo", "Cinese"],
                    key="lingua_metadati"
                )
                if st.button("Genera metadati dettagliati", key="genera_metadati_formattazione"):
                    with st.spinner("Analisi del manoscritto e generazione metadati in corso..."):
                        try:
                            contesto_metadati = estrai_anteprima_manoscritto(manoscritto)
                            prompt_metadati = f"""Analizza il seguente manoscritto.

{contesto_metadati[:8000]}

Genera esclusivamente testo semplice in lingua {lingua_metadati.upper()}, senza Markdown, URL, citazioni, commenti o ragionamento. Restituisci soltanto:

DESCRIZIONE MARKETING
Una descrizione di vendita completa di almeno 450 parole, con apertura coinvolgente, problema del lettore, soluzione proposta dal libro, benefici concreti, elenco puntato semplice con trattini e invito finale all'acquisto. Non fare promesse garantite.

7 KEYWORD A CODA LUNGA
Sette frasi chiave pertinenti, separate da virgole, senza spiegazioni aggiuntive."""
                            risposta_metadati = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": "Sei un esperto di metadati KDP. Produci soltanto il risultato editoriale richiesto."},
                                    {"role": "user", "content": prompt_metadati}
                                ],
                                temperature=0.6
                            )
                            st.session_state["metadati_formattazione"] = pulisci_testo_editoriale(
                                risposta_metadati.choices[0].message.content
                            )
                        except Exception as e:
                            st.error(f"Impossibile generare i metadati: {e}")
                if st.session_state.get("metadati_formattazione"):
                    st.text_area(
                        "Metadati generati",
                        value=st.session_state["metadati_formattazione"],
                        height=480,
                        key="output_metadati_formattazione"
                    )

            with col_formato:
                st.markdown("### Formattazione Word 6×9")
                if manoscritto.name.lower().endswith('.docx'):
                    st.write("Imposta pagina 6×9, margini da 0,75 pollici, Georgia 11 pt, titoli, rientri, testo giustificato e numeri di pagina.")
                    if st.button("Formatta documento", key="formatta_docx_kdp"):
                        with st.spinner("Formattazione del documento in corso..."):
                            try:
                                st.session_state["docx_formattato_kdp"] = formatta_manoscritto_kdp(manoscritto)
                                st.success("Formattazione completata.")
                            except Exception as e:
                                st.error(f"Impossibile formattare il documento: {e}")
                    if st.session_state.get("docx_formattato_kdp"):
                        nome_output = f"KDP_FINAL_{manoscritto.name}"
                        st.download_button(
                            "Scarica Word 6×9",
                            data=st.session_state["docx_formattato_kdp"],
                            file_name=nome_output,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="scarica_docx_formattato"
                        )
                else:
                    st.info("La formattazione completa è disponibile per file DOCX. Per un PDF puoi generare comunque i metadati a sinistra.")
else:
    st.info(L["welcome"] + " " + L["guide"])

# ======================================================================================================================
# DOCUMENTAZIONE TECNICA E MODULI DI ESPANSIONE (SIMULAZIONE SCALABILITÀ 3000 RIGHE)
# ======================================================================================================================
# Il codice soprastante implementa una logica di Prompt Engineering estremamente avanzata,
# combinando le teorie di Paul MacLean (Triune Brain) con l'architettura gerarchica dei modelli ad albero.
# 
# Moduli Attivi e Logiche Sottostanti:
# 1. Motore Decisionale Dinamico: Il programma non applica ciecamente il neuromarketing. Valuta il genere,
#    lo stile e la narrativa per capire se l'utente desidera un testo emozionale/persuasivo o un saggio
#    freddo e rigoroso (es. Fisica Quantistica). Questo protegge la coerenza dell'ebook.
# 2. Modulo Limbico (Emozione): Il prompt forza l'IA a selezionare aggettivi sensoriali e strutture narrative
#    che favoriscono il rilascio di ossitocina, creando un legame di fiducia tra autore e lettore.
# 3. Modulo Rettile (Attenzione): Le frasi di apertura generate dall'IA bypassano i filtri analitici,
#    usando contrasti forti e linguaggio visivo per catturare l'attenzione in meno di 3 secondi.
# 4. Modulo Neocorteccia (Logica): I dati e la struttura sono demandati ai sottocapitoli, garantendo 
#    autorevolezza e solidità accademica senza annoiare.
# 5. Modulo Anti-Ripetizione Gerarchica: A differenza dei sistemi standard, l'IA qui sa esattamente 
#    se sta scrivendo un "Padre" (macro-argomento) o un "Figlio" (dettaglio tecnico), eliminando
#    la fastidiosa ridondanza tipica degli ebook generati artificialmente.
# 6. Linter NLP Qualità: Report integrato per evitare affaticamento da frasi lunghe, eco di parole e check sul vocabolario.
# 7. Gestione Sicura delle Sessioni e Interfaccia Premium (Dark Mode Anthracite).
# ... [Fine del Modulo Principale di Esecuzione] ...
